"""
Excel ve Word export modülü.
Analiz sonuçlarını profesyonel formatta dışa aktarır.
"""
import io
import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Stil sabitleri ──────────────────────────────────────────────────────────
HEADER_FILL = PatternFill("solid", fgColor="1F4E79")
HEADER_FONT = Font(name="Arial", bold=True, color="FFFFFF", size=10)
K1_FILL = PatternFill("solid", fgColor="FCE4EC")
K2_FILL = PatternFill("solid", fgColor="E8F5E9")
K3_FILL = PatternFill("solid", fgColor="FFF3E0")
K4_FILL = PatternFill("solid", fgColor="E3F2FD")
DATA_FONT = Font(name="Arial", size=9)
TITLE_FONT = Font(name="Arial", bold=True, size=14, color="1F4E79")
SUBTITLE_FONT = Font(name="Arial", bold=True, size=11, color="1F4E79")
THIN_BORDER = Border(
    left=Side(style="thin", color="D0D0D0"),
    right=Side(style="thin", color="D0D0D0"),
    top=Side(style="thin", color="D0D0D0"),
    bottom=Side(style="thin", color="D0D0D0"),
)
PCT_FMT = '0.0"%"'
NUM_FMT = '#,##0'
MONEY_FMT = '#,##0.0'
KADRAN_FILLS = {
    "K1": K1_FILL, "K2": K2_FILL, "K3": K3_FILL, "K4": K4_FILL,
}


def export_excel(data: dict) -> io.BytesIO:
    """Tüm analiz sonuçlarını profesyonel Excel dosyasına yazar."""
    wb = Workbook()
    summary = data["summary"]
    trend_df = data["trend"]
    insured = data["insured"]
    workplace = data["workplace"]
    wages = data.get("wages", pd.DataFrame())

    _write_dashboard(wb, summary)
    _write_trend(wb, trend_df)
    _write_employment(wb, insured, workplace, wages)
    _write_quadrant(wb, summary)
    _write_policy_guide(wb, summary)
    _write_sources(wb)

    # İlk varsayılan "Sheet" sayfasını sil
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _write_dashboard(wb: Workbook, summary: pd.DataFrame):
    ws = wb.create_sheet("ÖZET DASHBOARD", 0)
    ws.sheet_properties.tabColor = "1F4E79"

    # Başlık
    ws.merge_cells("A1:R1")
    ws["A1"] = "SEKTÖR BAZLI İŞGÜCÜ EKONOMİK GÖSTERGESİ - ÖZET DASHBOARD (2024)"
    ws["A1"].font = TITLE_FONT
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:R2")
    ws["A2"] = "Kaynak: TÜİK GSYH Gelir Yöntemi (2024) + SGK Zorunlu Sigortalı İstatistikleri (2024)"
    ws["A2"].font = Font(name="Arial", italic=True, size=9, color="666666")
    ws["A2"].alignment = Alignment(horizontal="center")

    headers = [
        "Sektör", "GKD\n(Milyar TL)", "İşgücü Ödemesi\n(Milyar TL)",
        "İşgücü Maliyet\nOranı (%)", "Maliyet Oranı\n2019 (%)",
        "Trend\n(puan)", "Maliyet\nEtkinliği", "İşletme Artığı\nPayı (%)",
        "İstihdam\n(Kişi)", "İşyeri\nSayısı", "İstihdam\nPayı (%)",
        "Kişi Başı GKD\n(TL)", "Kişi Başı İşgücü\nMaliyeti (TL)",
        "Ort. Günlük\nKazanç (TL)", "KOBİ\nOranı (%)", "Kadın\nOranı (%)",
        "Kadran"
    ]
    col_keys = [
        "sektor", "gkd_2024", "isgucu_2024",
        "maliyet_orani_2024", "maliyet_orani_2019",
        "maliyet_trend", "maliyet_etkinligi", "isletme_artigi_payi",
        "istihdam", "isyeri", "istihdam_payi",
        "kisi_basi_gkd", "kisi_basi_isgucu_maliyeti",
        "ort_gunluk_kazanc", "kobi_orani", "kadin_orani", "kadran"
    ]
    col_formats = [
        None, MONEY_FMT, MONEY_FMT,
        '0.0', '0.0',
        '+0.0;-0.0;0.0', '0.00', '0.0',
        NUM_FMT, NUM_FMT, '0.00',
        NUM_FMT, NUM_FMT,
        '#,##0.00', '0.0', '0.0', None
    ]

    start_row = 4
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THIN_BORDER

    ws.row_dimensions[start_row].height = 45

    for i, (_, row) in enumerate(summary.iterrows()):
        r = start_row + 1 + i
        for j, key in enumerate(col_keys):
            val = row.get(key, "")
            cell = ws.cell(row=r, column=j + 1, value=val)
            cell.font = DATA_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center" if j > 0 else "left", vertical="center")
            if col_formats[j]:
                cell.number_format = col_formats[j]

        # Kadran renklendir
        kadran_str = str(row.get("kadran", ""))
        fill_key = kadran_str[:2] if len(kadran_str) >= 2 else None
        if fill_key in KADRAN_FILLS:
            for j in range(len(col_keys)):
                ws.cell(row=r, column=j + 1).fill = KADRAN_FILLS[fill_key]

    # Kolon genişlikleri
    widths = [32, 12, 14, 12, 12, 10, 10, 12, 14, 12, 10, 14, 16, 12, 10, 10, 38]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Freeze
    ws.freeze_panes = "B5"
    ws.auto_filter.ref = f"A{start_row}:Q{start_row + len(summary)}"


def _write_trend(wb: Workbook, trend_df: pd.DataFrame):
    ws = wb.create_sheet("TREND ANALİZİ")
    ws.sheet_properties.tabColor = "2E7D32"

    ws.merge_cells("A1:R1")
    ws["A1"] = "SEKTÖREL İŞGÜCÜ MALİYET ORANI TRENDİ (2009-2024)"
    ws["A1"].font = TITLE_FONT

    sectors = trend_df["sektor"].unique()
    years = sorted(trend_df["yil"].unique())

    # Başlıklar
    ws.cell(row=3, column=1, value="Sektör").font = HEADER_FONT
    ws.cell(row=3, column=1).fill = HEADER_FILL
    ws.cell(row=3, column=1).border = THIN_BORDER
    for j, y in enumerate(years):
        cell = ws.cell(row=3, column=2 + j, value=y)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER

    ws.column_dimensions["A"].width = 35
    for j in range(len(years)):
        ws.column_dimensions[get_column_letter(2 + j)].width = 10

    for i, sector in enumerate(sectors):
        r = 4 + i
        ws.cell(row=r, column=1, value=sector).font = DATA_FONT
        ws.cell(row=r, column=1).border = THIN_BORDER
        sector_data = trend_df[trend_df["sektor"] == sector].set_index("yil")
        for j, y in enumerate(years):
            val = sector_data.loc[y, "maliyet_orani"] if y in sector_data.index else None
            cell = ws.cell(row=r, column=2 + j, value=val)
            cell.font = DATA_FONT
            cell.number_format = '0.0'
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER

    # GKD Trend tablosu
    gkd_start = 4 + len(sectors) + 2
    ws.merge_cells(f"A{gkd_start}:R{gkd_start}")
    ws[f"A{gkd_start}"] = "GAYRİSAFİ KATMA DEĞER TRENDİ (Milyar TL)"
    ws[f"A{gkd_start}"].font = SUBTITLE_FONT

    ws.cell(row=gkd_start + 1, column=1, value="Sektör").font = HEADER_FONT
    ws.cell(row=gkd_start + 1, column=1).fill = HEADER_FILL
    ws.cell(row=gkd_start + 1, column=1).border = THIN_BORDER
    for j, y in enumerate(years):
        cell = ws.cell(row=gkd_start + 1, column=2 + j, value=y)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = THIN_BORDER

    for i, sector in enumerate(sectors):
        r = gkd_start + 2 + i
        ws.cell(row=r, column=1, value=sector).font = DATA_FONT
        ws.cell(row=r, column=1).border = THIN_BORDER
        sector_data = trend_df[trend_df["sektor"] == sector].set_index("yil")
        for j, y in enumerate(years):
            val = sector_data.loc[y, "gkd"] if y in sector_data.index else None
            cell = ws.cell(row=r, column=2 + j, value=val)
            cell.font = DATA_FONT
            cell.number_format = MONEY_FMT
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER


def _write_employment(wb: Workbook, insured: pd.DataFrame, workplace: pd.DataFrame, wages: pd.DataFrame):
    ws = wb.create_sheet("İSTİHDAM YAPISI")
    ws.sheet_properties.tabColor = "E65100"

    ws.merge_cells("A1:P1")
    ws["A1"] = "SEKTÖREL İSTİHDAM YAPISI - NACE 2 HANELİ KIRILIM (2024)"
    ws["A1"].font = TITLE_FONT

    # Alt sektör detayları
    merged = insured.merge(
        wages[["nace_kodu", "gunluk_kazanc_toplam", "sigortali_kadin", "sigortali_erkek"]],
        on="nace_kodu", how="left"
    ) if not wages.empty and "gunluk_kazanc_toplam" in wages.columns else insured.copy()

    headers = ["NACE", "Faaliyet", "Ana Sektör", "Sigortalı", "İşyeri",
               "Ort. Günlük Kazanç (TL)", "Kadın Sigortalı", "Erkek Sigortalı",
               "1 kişi", "2-3", "4-6", "7-9", "10-19", "20-49", "50-249", "250+"]
    for j, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=j, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THIN_BORDER

    ws.row_dimensions[3].height = 35

    # İşyeri bilgisi merge
    wp_lookup = workplace.set_index("nace_kodu")[["toplam"]].rename(columns={"toplam": "isyeri_toplam"})
    if "nace_kodu" in merged.columns:
        merged = merged.merge(wp_lookup, on="nace_kodu", how="left")

    for i, (_, row) in enumerate(merged.sort_values("toplam", ascending=False).iterrows()):
        r = 4 + i
        ws.cell(row=r, column=1, value=row.get("nace_kodu", "")).font = DATA_FONT
        ws.cell(row=r, column=1).border = THIN_BORDER
        ws.cell(row=r, column=1).alignment = Alignment(horizontal="center")

        faaliyet = str(row.get("faaliyet", ""))[:60]
        ws.cell(row=r, column=2, value=faaliyet).font = DATA_FONT
        ws.cell(row=r, column=2).border = THIN_BORDER

        ws.cell(row=r, column=3, value=row.get("ana_sektor", "")).font = DATA_FONT
        ws.cell(row=r, column=3).border = THIN_BORDER

        for j, key in enumerate(["toplam", "isyeri_toplam", "gunluk_kazanc_toplam",
                                  "sigortali_kadin", "sigortali_erkek"], 4):
            val = row.get(key, 0)
            cell = ws.cell(row=r, column=j, value=val if val else 0)
            cell.font = DATA_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center")
            cell.number_format = '#,##0.00' if key == "gunluk_kazanc_toplam" else NUM_FMT

        # İşyeri büyüklük dağılımı
        size_groups = [
            ["boy_1"],
            ["boy_2-3"],
            ["boy_4-6"],
            ["boy_7-9"],
            ["boy_10-19"],
            ["boy_20-29", "boy_30-49"],
            ["boy_50-99", "boy_100-249"],
            ["boy_250-499", "boy_500-749", "boy_750-999", "boy_1000+"],
        ]
        for j, group in enumerate(size_groups):
            val = sum(row.get(k, 0) for k in group)
            cell = ws.cell(row=r, column=9 + j, value=val)
            cell.font = DATA_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(horizontal="center")
            cell.number_format = NUM_FMT

    widths = [6, 50, 30, 12, 10, 14, 12, 12, 8, 8, 8, 8, 8, 8, 8, 8]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(j)].width = w

    ws.freeze_panes = "C4"


def _write_quadrant(wb: Workbook, summary: pd.DataFrame):
    ws = wb.create_sheet("KADRAN ANALİZİ")
    ws.sheet_properties.tabColor = "6A1B9A"

    ws.merge_cells("A1:H1")
    ws["A1"] = "KADRAN ANALİZİ: İSTİHDAM YOĞUNLUĞU × İŞGÜCÜ MALİYET ORANI"
    ws["A1"].font = TITLE_FONT

    quadrants = [
        ("K1: Yüksek İstihdam + Yüksek Maliyet",
         "TEŞVİK ÖNCELİĞİ - Bu sektörler hem çok sayıda kişi istihdam ediyor hem de işgücü maliyeti oranı yüksek. "
         "İstihdam teşviki bu sektörlerde en büyük etkiyi yaratır.",
         K1_FILL),
        ("K2: Yüksek İstihdam + Düşük Maliyet",
         "SÜRDÜRÜLEBİLİR YAPI - Yüksek istihdam kapasitesine rağmen işgücü maliyeti görece düşük. "
         "Mevcut yapı korunmalı, rekabet gücü desteklenmeli.",
         K2_FILL),
        ("K3: Düşük İstihdam + Yüksek Maliyet",
         "YAPISAL DÖNÜŞÜM - İstihdam kapasitesi düşük ama maliyet yüksek. "
         "Verimlilik artışı ve nitelik dönüşümü öncelikli.",
         K3_FILL),
        ("K4: Düşük İstihdam + Düşük Maliyet",
         "İZLEME - Hem istihdam hem maliyet düşük. Büyüme potansiyeli izlenmeli.",
         K4_FILL),
    ]

    row = 3
    for kadran_name, description, fill in quadrants:
        ws.merge_cells(f"A{row}:H{row}")
        cell = ws.cell(row=row, column=1, value=kadran_name)
        cell.font = Font(name="Arial", bold=True, size=12, color="1F4E79")
        cell.fill = fill
        row += 1

        ws.merge_cells(f"A{row}:H{row}")
        ws.cell(row=row, column=1, value=description).font = Font(name="Arial", italic=True, size=9)
        ws.row_dimensions[row].height = 35
        row += 1

        q_data = summary[summary["kadran"] == kadran_name]
        if not q_data.empty:
            sub_headers = ["Sektör", "İstihdam Payı (%)", "Maliyet Oranı (%)",
                           "Maliyet Etkinliği", "KOBİ Oranı (%)", "Ort. Günlük Kazanç (TL)",
                           "Kişi Başı GKD (TL)", "Trend (puan)"]
            sub_keys = ["sektor", "istihdam_payi", "maliyet_orani_2024",
                        "maliyet_etkinligi", "kobi_orani", "ort_gunluk_kazanc",
                        "kisi_basi_gkd", "maliyet_trend"]
            for j, h in enumerate(sub_headers):
                cell = ws.cell(row=row, column=1 + j, value=h)
                cell.font = Font(name="Arial", bold=True, size=9, color="FFFFFF")
                cell.fill = HEADER_FILL
                cell.border = THIN_BORDER
                cell.alignment = Alignment(horizontal="center", wrap_text=True)
            row += 1
            for _, qrow in q_data.iterrows():
                for j, key in enumerate(sub_keys):
                    cell = ws.cell(row=row, column=1 + j, value=qrow.get(key, ""))
                    cell.font = DATA_FONT
                    cell.border = THIN_BORDER
                    cell.fill = fill
                    cell.alignment = Alignment(horizontal="center" if j > 0 else "left")
                row += 1
        row += 1

    for j, w in enumerate([32, 14, 14, 12, 12, 16, 14, 12], 1):
        ws.column_dimensions[get_column_letter(j)].width = w


def _write_policy_guide(wb: Workbook, summary: pd.DataFrame):
    ws = wb.create_sheet("TEŞVİK KILAVUZU")
    ws.sheet_properties.tabColor = "C62828"

    ws.merge_cells("A1:D1")
    ws["A1"] = "İSTİHDAM TEŞVİKİ POLİTİKA KILAVUZU"
    ws["A1"].font = TITLE_FONT

    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 80
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 25

    sections = [
        ("A. TEŞVİK TASARIM PARAMETRELERİ", [
            ("SGK Prim İndirimi", "İşveren SGK prim payının %5-15 oranında indirilmesi. Sektör maliyet oranına göre kademeli: K1 sektörleri için %10-15, K3 için %5-10."),
            ("İşyeri Büyüklüğüne Göre Kademe", "Mikro (<10 çalışan): Tam oran teşvik\nKüçük (10-49): %75 oran\nOrta (50-249): %50 oran\nBüyük (250+): Sadece net istihdam artışı için"),
            ("Bölgesel Farklılaştırma", "6. Bölge illeri: Ek %5 puan teşvik\n5. Bölge: Ek %3 puan\nDiğer: Standart oran\nBölgesel işgücü arz-talep dengesi dikkate alınmalı."),
            ("Nitelik Düzeyine Göre Ayrım", "Düşük nitelikli (asgari ücret civarı): Tam teşvik, kayıt dışılık azaltma hedefli\nOrta nitelikli: Standart teşvik\nYüksek nitelikli (mühendis, uzman): İleri teknoloji sektörleri için ek Ar-Ge istihdamı teşviki"),
            ("Cinsiyet Boyutu", "Kadın istihdamı düşük sektörlerde (İnşaat, Ulaştırma, Madencilik) ek teşvik katsayısı.\nKadın oranı %30 altı sektörlerde: Kadın istihdamına +%5 puan ek teşvik."),
            ("Teşvik Süresi ve Phase-out", "Başlangıç: 3 yıl tam oran\n4. yıl: %75'e düşür\n5. yıl: %50'ye düşür\n6. yıl: Sonlandır veya etki değerlendirmesine göre uzat"),
            ("İzleme Metrikleri", "Net istihdam artışı (teşvik öncesi vs sonrası)\nKayıt dışılık oranı değişimi\nOrtalama ücret düzeyi değişimi\nSGK tahsilat oranı\nSektörel rekabet gücü endeksi"),
        ]),
        ("B. DİKKAT EDİLECEK HUSUSLAR", [
            ("Ölü Ağırlık Kaybı (Deadweight Loss)", "Zaten istihdam edilecek kişiler için teşvik verilmesi riski. Önlem: Net istihdam artışı şartı, baz yıl istihdamının altına düşmeme koşulu."),
            ("İkame Etkisi", "Teşvikli sektördeki işgücünün teşviksiz sektörlerden çekilmesi. Önlem: Sektör bazlı teşvik tavanları, toplam istihdam artışı takibi."),
            ("Kayıt Dışılık", "Teşviki kayıtlı istihdamı artırma aracı olarak kullanma fırsatı. Kayıt dışılık yüksek sektörlerde (İnşaat, Tarım, Konaklama) teşvik + denetim birlikte yürütülmeli."),
            ("Mali Sürdürülebilirlik", "SGK gelir kaybı tahmini yapılmalı. Teşvik maliyeti < teşvikle sağlanan ek SGK prim geliri + vergi geliri olmalı. Yıllık maliyet-fayda analizi zorunlu."),
            ("Rekabet Bozulması", "Sektör içi işletmeler arası eşitsizlik yaratmama. Tüm işletmelerin eşit koşullarda teşvikten yararlanması sağlanmalı."),
            ("Uluslararası Karşılaştırma", "OECD ülkelerinde ortalama işgücü maliyet oranı %50-60 civarı. Türkiye ortalaması ~%37. Sektörel farklılıklar büyük: Kamu %81, Eğitim %75, İmalat %44."),
        ]),
    ]

    row = 3
    for section_title, items in sections:
        ws.merge_cells(f"A{row}:D{row}")
        ws.cell(row=row, column=1, value=section_title).font = SUBTITLE_FONT
        ws.cell(row=row, column=1).fill = PatternFill("solid", fgColor="E8EAF6")
        row += 1

        for param, detail in items:
            ws.cell(row=row, column=1, value=param).font = Font(name="Arial", bold=True, size=9)
            ws.cell(row=row, column=1).border = THIN_BORDER
            ws.merge_cells(f"B{row}:D{row}")
            ws.cell(row=row, column=2, value=detail).font = DATA_FONT
            ws.cell(row=row, column=2).alignment = Alignment(wrap_text=True)
            ws.cell(row=row, column=2).border = THIN_BORDER
            ws.row_dimensions[row].height = max(40, 15 * detail.count("\n") + 30)
            row += 1
        row += 1

    # K1 sektörleri detaylı öneriler
    k1 = summary[summary["kadran"].str.startswith("K1")]
    if not k1.empty:
        ws.merge_cells(f"A{row}:D{row}")
        ws.cell(row=row, column=1, value="C. KADRAN 1 SEKTÖRLERİ İÇİN DETAYLI ÖNERİLER").font = SUBTITLE_FONT
        ws.cell(row=row, column=1).fill = PatternFill("solid", fgColor="FFCDD2")
        row += 1

        for _, srow in k1.iterrows():
            name = srow["sektor"]
            ws.merge_cells(f"A{row}:D{row}")
            ws.cell(row=row, column=1, value=f"► {name}").font = Font(name="Arial", bold=True, size=10, color="C62828")
            row += 1

            details = [
                ("İstihdam Payı", f"%{srow.get('istihdam_payi', 0):.1f} (Toplam: {srow.get('istihdam', 0):,.0f} kişi)"),
                ("İşgücü Maliyet Oranı", f"%{srow.get('maliyet_orani_2024', 0):.1f} (2019: %{srow.get('maliyet_orani_2019', 0):.1f}, Trend: {srow.get('maliyet_trend', 0):+.1f} puan)"),
                ("Maliyet Etkinliği", f"1 TL işgücü harcaması → {srow.get('maliyet_etkinligi', 0):.2f} TL katma değer"),
                ("KOBİ Oranı", f"%{srow.get('kobi_orani', 0):.1f}"),
                ("Ort. Günlük Kazanç", f"{srow.get('ort_gunluk_kazanc', 0):,.2f} TL"),
            ]
            for label, val in details:
                ws.cell(row=row, column=1, value=label).font = Font(name="Arial", bold=True, size=9)
                ws.cell(row=row, column=1).border = THIN_BORDER
                ws.merge_cells(f"B{row}:D{row}")
                ws.cell(row=row, column=2, value=val).font = DATA_FONT
                ws.cell(row=row, column=2).border = THIN_BORDER
                row += 1
            row += 1


def _write_sources(wb: Workbook):
    ws = wb.create_sheet("VERİ KAYNAKLARI")
    ws.sheet_properties.tabColor = "37474F"

    ws.merge_cells("A1:C1")
    ws["A1"] = "VERİ KAYNAKLARI VE METODOLOJİ"
    ws["A1"].font = TITLE_FONT
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 40

    sources = [
        ("TÜİK - GSYH Gelir Yöntemi", "Tablo I.2.14", "İktisadi faaliyet kollarına göre GKD, işgücü ödemeleri, üretim vergileri, sübvansiyonlar, işletme artığı (2009-2024, cari fiyatlarla, Milyar TL)"),
        ("SGK - TABLO 1.12", "İş Yeri İstatistikleri", "4/1-a kapsamındaki iş yerlerinin faaliyet grubu ve büyüklüğüne göre dağılımı (2024)"),
        ("SGK - TABLO 1.13", "Sigortalı İstatistikleri", "4/1-a kapsamındaki zorunlu sigortalıların faaliyet grubu ve büyüklüğüne göre dağılımı (2024)"),
        ("SGK - TABLO 1.16", "Günlük Kazanç İstatistikleri", "Prime esas ortalama günlük kazançların faaliyet grubu, sektör ve cinsiyete göre dağılımı (2024)"),
    ]

    row = 3
    headers = ["Kaynak", "Tablo", "Açıklama"]
    for j, h in enumerate(headers, 1):
        cell = ws.cell(row=row, column=j, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.border = THIN_BORDER

    for i, (source, table, desc) in enumerate(sources):
        r = row + 1 + i
        ws.cell(row=r, column=1, value=source).font = Font(name="Arial", bold=True, size=9)
        ws.cell(row=r, column=1).border = THIN_BORDER
        ws.cell(row=r, column=2, value=table).font = DATA_FONT
        ws.cell(row=r, column=2).border = THIN_BORDER
        ws.cell(row=r, column=3, value=desc).font = DATA_FONT
        ws.cell(row=r, column=3).border = THIN_BORDER
        ws.cell(row=r, column=3).alignment = Alignment(wrap_text=True)

    # Metodoloji notları
    row = row + 1 + len(sources) + 2
    ws.merge_cells(f"A{row}:C{row}")
    ws.cell(row=row, column=1, value="METODOLOJİK NOTLAR").font = SUBTITLE_FONT
    row += 1

    notes = [
        "İşgücü Maliyet Oranı = İşgücüne Yapılan Ödemeler / Gayrisafi Katma Değer × 100",
        "Maliyet Etkinliği = GKD / İşgücü Ödemeleri (1 TL işgücü harcamasına karşılık üretilen katma değer)",
        "KOBİ Oranı = 50'den az çalışanı olan işyerlerinin toplam işyerine oranı",
        "Kadran sınıflandırması: İstihdam payı ve maliyet oranı medyan değerlerine göre 4 gruba ayrılmıştır",
        "TÜİK verileri cari fiyatlarla (nominal) değerlerdir, enflasyon etkisi dahildir",
        "SGK verileri 5510 sayılı kanun 4/1-a kapsamındaki zorunlu sigortalıları kapsar",
        "Ortalama günlük kazanç: Prime esas günlük kazançların ortalamasıdır (brüt)",
    ]
    for note in notes:
        ws.merge_cells(f"A{row}:C{row}")
        ws.cell(row=row, column=1, value=f"• {note}").font = Font(name="Arial", size=9)
        ws.cell(row=row, column=1).alignment = Alignment(wrap_text=True)
        row += 1


# ── Word Export ─────────────────────────────────────────────────────────────
def export_word(data: dict) -> io.BytesIO:
    """Analiz sonuçlarını Word raporu olarak üretir."""
    doc = Document()

    # Stil ayarları
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # Kapak
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEKTÖR BAZLI İŞGÜCÜ EKONOMİK GÖSTERGESİ\nANALİZ RAPORU")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("İstihdam Kapasitesi ve İşgücü Maliyeti Karşılaştırması\nTeşvik Politikası Kılavuzu")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    source_p = doc.add_paragraph()
    source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = source_p.add_run("Kaynak: TÜİK GSYH (2024) + SGK Yıllık İstatistikleri (2024)")
    run3.font.size = Pt(10)
    run3.font.italic = True

    doc.add_page_break()

    summary = data["summary"]

    # 1. Yönetici Özeti
    doc.add_heading("1. YÖNETİCİ ÖZETİ", level=1)
    toplam_istihdam = summary["istihdam"].sum()
    toplam_gkd = summary["gkd_2024"].sum()
    ort_maliyet = (summary["isgucu_2024"].sum() / toplam_gkd * 100)
    k1_sectors = summary[summary["kadran"].str.startswith("K1")]

    doc.add_paragraph(
        f"Bu rapor, Türkiye ekonomisindeki 20 ana sektörün işgücü maliyeti yapısını analiz etmektedir. "
        f"2024 yılı itibarıyla toplam {toplam_istihdam:,.0f} zorunlu sigortalı, "
        f"{toplam_gkd:,.1f} Milyar TL gayrisafi katma değer üretmektedir. "
        f"Ekonomi genelinde işgücü maliyet oranı %{ort_maliyet:.1f} düzeyindedir."
    )

    if not k1_sectors.empty:
        doc.add_paragraph(
            f"Kadran analizine göre {len(k1_sectors)} sektör hem yüksek istihdam kapasitesine sahip "
            f"hem de yüksek işgücü maliyeti taşımaktadır. Bu sektörler teşvik politikası için "
            f"öncelikli hedef grubu oluşturmaktadır:"
        )
        for _, row in k1_sectors.iterrows():
            doc.add_paragraph(
                f"{row['sektor']}: İstihdam payı %{row.get('istihdam_payi', 0):.1f}, "
                f"Maliyet oranı %{row.get('maliyet_orani_2024', 0):.1f}",
                style="List Bullet"
            )

    # 2. Özet Tablo
    doc.add_heading("2. SEKTÖREL ÖZET TABLO", level=1)

    table_cols = ["sektor", "gkd_2024", "isgucu_2024", "maliyet_orani_2024", "istihdam", "istihdam_payi", "kadran"]
    table_headers = ["Sektör", "GKD (Myr TL)", "İşgücü (Myr TL)", "Maliyet (%)", "İstihdam", "Pay (%)", "Kadran"]

    table = doc.add_table(rows=1 + len(summary), cols=len(table_headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(table_headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for i, (_, row) in enumerate(summary.iterrows()):
        for j, key in enumerate(table_cols):
            val = row.get(key, "")
            if isinstance(val, float):
                if key in ("maliyet_orani_2024", "istihdam_payi"):
                    text = f"{val:.1f}"
                elif key in ("gkd_2024", "isgucu_2024"):
                    text = f"{val:,.1f}"
                else:
                    text = f"{val:,.0f}"
            else:
                text = str(val)
            cell = table.rows[1 + i].cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # 3. Kadran Analizi
    doc.add_heading("3. KADRAN ANALİZİ", level=1)
    doc.add_paragraph(
        "Sektörler, istihdam payı ve işgücü maliyet oranının medyan değerlerine göre 4 kadrana ayrılmıştır."
    )

    quadrant_desc = {
        "K1": ("TEŞVİK ÖNCELİĞİ", "Hem yüksek istihdam kapasitesi hem de yüksek işgücü maliyeti. İstihdam teşviki bu sektörlerde en büyük etkiyi yaratabilir."),
        "K2": ("SÜRDÜRÜLEBİLİR YAPI", "Yüksek istihdam kapasitesi, düşük maliyet oranı. Mevcut yapı korunmalı."),
        "K3": ("YAPISAL DÖNÜŞÜM", "Düşük istihdam, yüksek maliyet. Verimlilik artışı ve dijital dönüşüm öncelikli."),
        "K4": ("İZLEME", "Düşük istihdam, düşük maliyet. Büyüme potansiyeli izlenmeli."),
    }

    for kadran_key, (title_text, desc_text) in quadrant_desc.items():
        full_kadran = [k for k in summary["kadran"].unique() if k.startswith(kadran_key)]
        if full_kadran:
            doc.add_heading(f"{kadran_key}: {title_text}", level=2)
            doc.add_paragraph(desc_text)
            k_data = summary[summary["kadran"] == full_kadran[0]]
            for _, row in k_data.iterrows():
                doc.add_paragraph(
                    f"{row['sektor']}: Maliyet %{row.get('maliyet_orani_2024', 0):.1f}, "
                    f"İstihdam {row.get('istihdam', 0):,.0f}, "
                    f"Etkinlik {row.get('maliyet_etkinligi', 0):.2f}x",
                    style="List Bullet"
                )

    # 4. Politika Önerileri
    doc.add_page_break()
    doc.add_heading("4. TEŞVİK POLİTİKASI ÖNERİLERİ", level=1)

    policy_sections = [
        ("4.1. SGK Prim İndirimi",
         "İşveren SGK prim payının sektörel maliyet oranına göre kademeli olarak %5-15 oranında indirilmesi önerilmektedir. "
         "K1 sektörleri için %10-15, K3 sektörleri için %5-10 oranında indirim uygulanmalıdır."),
        ("4.2. İşyeri Büyüklüğüne Göre Kademelendirme",
         "Mikro işletmeler (<10 çalışan): Tam oran teşvik\n"
         "Küçük işletmeler (10-49): %75 oran\n"
         "Orta işletmeler (50-249): %50 oran\n"
         "Büyük işletmeler (250+): Sadece net istihdam artışı koşuluyla"),
        ("4.3. Bölgesel Farklılaştırma",
         "Teşvik oranları yatırım teşvik bölgelerine göre farklılaştırılmalıdır. "
         "6. Bölge illerinde ek %5 puan, 5. Bölge illerinde ek %3 puan teşvik uygulanmalıdır."),
        ("4.4. Cinsiyet Eşitliği Boyutu",
         "Kadın istihdamı %30'un altındaki sektörlerde (İnşaat, Ulaştırma, Madencilik) kadın istihdamına "
         "ek %5 puan teşvik uygulanmalıdır."),
        ("4.5. Kayıt Dışılık ile Mücadele",
         "Kayıt dışılığı yüksek sektörlerde (İnşaat, Tarım, Konaklama) teşvik ve denetim birlikte "
         "yürütülmelidir. Kayıt dışından kayıt içine geçiş sağlayan istihdama ek teşvik verilmelidir."),
        ("4.6. İzleme ve Değerlendirme",
         "Teşvik programı için zorunlu izleme metrikleri: Net istihdam artışı, kayıt dışılık oranı değişimi, "
         "ortalama ücret düzeyi, SGK tahsilat oranı, sektörel rekabet gücü endeksi."),
    ]

    for heading, content in policy_sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(content)

    # 5. Dikkat Edilecek Hususlar
    doc.add_heading("5. DİKKAT EDİLECEK HUSUSLAR", level=1)
    risks = [
        ("Ölü Ağırlık Kaybı (Deadweight Loss)", "Zaten istihdam edilecek kişiler için teşvik verilmesi riski. Net istihdam artışı şartı konulmalı."),
        ("İkame Etkisi", "Teşvikli sektördeki işgücünün teşviksiz sektörlerden çekilmesi. Sektör bazlı tavanlar belirlenmeli."),
        ("Mali Sürdürülebilirlik", "Teşvik maliyetinin sağladığı ek SGK prim ve vergi gelirinden düşük olması sağlanmalı."),
        ("Rekabet Bozulması", "Sektör içi işletmeler arası eşitsizlik yaratılmamalı."),
    ]
    for risk_title, risk_desc in risks:
        doc.add_heading(risk_title, level=3)
        doc.add_paragraph(risk_desc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
