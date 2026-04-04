"""
Küresel Ekonomik Gelişmeler Işığında Türkiye'de Sektörel İstihdam Teşviki:
Akademik Rapor Oluşturucu
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import sys
from datetime import date

REPORT_VERSION = "2.4"
REPORT_DATE = date.today().strftime("%d %B %Y").replace(
    "January", "Ocak").replace("February", "Şubat").replace(
    "March", "Mart").replace("April", "Nisan").replace(
    "May", "Mayıs").replace("June", "Haziran").replace(
    "July", "Temmuz").replace("August", "Ağustos").replace(
    "September", "Eylül").replace("October", "Ekim").replace(
    "November", "Kasım").replace("December", "Aralık")

sys.path.insert(0, os.path.dirname(__file__))
from data_loader import load_all_data


def create_academic_report(data: dict, output_path: str):
    doc = Document()

    # ── Stil ayarları ───────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Cm(1.25)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Heading stilleri
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        hs.paragraph_format.first_line_indent = Cm(0)

    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)

    def add_para(text, bold=False, italic=False, align=None, indent=True):
        p = doc.add_paragraph()
        if not indent:
            p.paragraph_format.first_line_indent = Cm(0)
        if align:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    summary = data["summary"]

    # ══════════════════════════════════════════════════════════════════════════
    # KAPAK
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Cm(0)
    run = title_p.add_run(
        "KÜRESEL EKONOMİK GELİŞMELER VE ULUSLARARASI KURUMLARIN\n"
        "SEKTÖREL TAHMİNLERİ IŞIĞINDA TÜRKİYE'DE\n"
        "DESTEKLENMESI GEREKEN SEKTÖRLER VE\n"
        "İSTİHDAM TEŞVİKİ POLİTİKA ÖNERİLERİ"
    )
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.first_line_indent = Cm(0)
    run2 = subtitle_p.add_run("Sektörel Analiz Raporu")
    run2.font.size = Pt(14)
    run2.italic = True

    doc.add_paragraph()

    source_p = doc.add_paragraph()
    source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_p.paragraph_format.first_line_indent = Cm(0)
    run3 = source_p.add_run(
        "Veri Kaynakları: TÜİK GSYH Gelir Yöntemi (2009-2024),\n"
        "SGK Yıllık İstatistikleri (2024),\n"
        "IMF, OECD, Dünya Bankası, ILO, UNIDO Raporları"
    )
    run3.font.size = Pt(11)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.first_line_indent = Cm(0)
    date_p.add_run(f"{REPORT_DATE}  |  Sürüm {REPORT_VERSION}").font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # İÇİNDEKİLER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("İÇİNDEKİLER", level=1)
    toc_items = [
        "1. Yönetici Özeti",
        "2. Giriş ve Amaç",
        "3. Küresel Ekonomik Görünüm ve Sektörel Trendler",
        "   3.1. Makroekonomik Çerçeve",
        "   3.2. Sektörel Dönüşüm Dinamikleri",
        "   3.3. İşgücü Piyasası Küresel Eğilimler",
        "4. Türkiye Ekonomisi: Sektörel Yapı Analizi",
        "   4.1. Gayrisafi Katma Değer ve İşgücü Maliyeti",
        "   4.2. İstihdam Yapısı ve Verimlilik",
        "   4.3. Kadran Analizi: Sektörel Sınıflandırma",
        "5. Uluslararası Kurum Değerlendirmeleri",
        "   5.1. IMF Perspektifi",
        "   5.2. OECD Önerileri",
        "   5.3. Dünya Bankası Değerlendirmesi",
        "   5.4. ILO ve UNIDO Bulgular",
        "6. Stratejik Öncelikli Sektörler",
        "   6.1. Desteklenmesi Gereken Sektörler ve Gerekçeleri",
        "   6.2. Sektörel Risk ve Fırsat Matrisi",
        "7. İstihdam Teşviki Politika Önerileri",
        "   7.1. Uluslararası Deneyimler ve Vergi Takozu Karşılaştırması",
        "   7.2. Teşvik Tasarım Parametreleri",
        "   7.3. Sektör Bazlı Teşvik Modelleri",
        "   7.4. AB Yeşil Mutabakat ve CBAM Uyumu",
        "   7.5. Dijital Dönüşüm ve Yapay Zeka",
        "   7.6. Beceri Açığı ve Aktif İşgücü Programlarının Rolü",
        "8. Sonuç ve Değerlendirme",
        "Kaynakça",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. YÖNETİCİ ÖZETİ
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. YÖNETİCİ ÖZETİ", level=1)

    toplam_istihdam = summary["istihdam"].sum()
    toplam_gkd = summary["gkd_2024"].sum()
    ort_maliyet = summary["isgucu_2024"].sum() / toplam_gkd * 100
    k1 = summary[summary["kadran"].str.startswith("K1")]

    add_para(
        f"Bu rapor, Türkiye ekonomisindeki 20 ana sektörün işgücü maliyeti yapısını, istihdam kapasitesini "
        f"ve verimlilik göstergelerini uluslararası karşılaştırmalı bir perspektiften analiz etmektedir. "
        f"2024 yılı itibarıyla 4/1-a kapsamında toplam {toplam_istihdam:,.0f} zorunlu sigortalı, "
        f"{toplam_gkd:,.1f} milyar TL gayrisafi katma değer üretmektedir. Ekonomi genelinde işgücü "
        f"maliyet oranı %{ort_maliyet:.1f} düzeyindedir."
    )

    add_para(
        "IMF'nin Ocak 2026 Dünya Ekonomik Görünümü güncellemesine göre küresel büyüme 2026'da %3,3 olarak "
        "öngörülmektedir (IMF, 2026). OECD, Türkiye'nin 2026'da %3,9 büyüyeceğini tahmin etmekte; ancak "
        "verimlilik artışının hızlandırılması, kadın işgücü katılımının yükseltilmesi ve hizmetler sektöründe "
        "düzenleyici reformlara ihtiyaç duyulduğunu vurgulamaktadır (OECD, 2025). Dünya Bankası ise "
        "Türkiye'nin 2027'de %4,4'e ulaşabileceğini ancak bunun yapısal reformlara bağlı olduğunu "
        "belirtmektedir (World Bank, 2025)."
    )

    add_para(
        f"Kadran analizine göre {len(k1)} sektör hem yüksek istihdam kapasitesine sahip hem de yüksek "
        f"işgücü maliyeti taşımaktadır. Bu sektörler —İmalat Sanayi, İnşaat, İdari ve Destek Hizmetleri, "
        f"Konaklama ve Yiyecek, Sağlık, Mesleki Faaliyetler ve Eğitim— toplam istihdamın yaklaşık "
        f"%70'ini oluşturmakta ve istihdam teşviki politikalarının öncelikli hedef grubunu "
        f"oluşturmaktadır. Kamu Yönetimi ve Savunma sektörü istatistiksel olarak Kadran 1'de yer "
        f"almakla birlikte, bu sektörün gayrisafi katma değeri ulusal hesaplarda piyasa çıktısı "
        f"yerine girdi maliyeti (ağırlıklı olarak ücretler) esas alınarak ölçüldüğünden maliyet "
        f"etkinliği oranı yapısal olarak ~1'e yakın çıkmaktadır. Piyasa mekanizması işlemediğinden "
        f"bu sektör kadran bazlı teşvik politikası kapsamı dışında tutulmaktadır."
    )

    add_para(
        "Rapor, AB Karbon Sınır Düzenleme Mekanizması (CBAM), yeşil dönüşüm, dijital dönüşüm ve yapay zeka "
        "gibi küresel megatrendleri Türkiye'nin sektörel yapısıyla ilişkilendirerek, desteklenmesi gereken "
        "sektörleri ve teşvik politikası parametrelerini ortaya koymaktadır."
    )

    add_para(
        "Analitik çerçeve olarak benimsenen kadran yöntemi, 20 sektörü iki eksen üzerinde —istihdam payı "
        "ve işgücü maliyet oranı— medyan değerlere göre dört gruba ayırmaktadır. Teşvik önceliği olan K1 "
        "sektörleri, hem yüksek istihdam kapasitesi hem de yüksek işgücü maliyet baskısı taşıdığından "
        "marjinal politika müdahalesinin en yüksek etkiyi yaratacağı gruptur. Bu sınıflandırma, kaynakların "
        "doğru sektörlere yönlendirilmesi ve teşvik tasarımının rasyonelleştirilmesi bakımından önemli bir "
        "analitik katkı sunmaktadır."
    )

    add_para(
        "İmalat sektörünün desteklenmesi, salt bu sektörün ötesine geçen çarpan etkileri nedeniyle stratejik "
        "öneme sahiptir. UNIDO Sanayi Kalkınma Raporu 2024, her imalat istihdamının ekonominin diğer "
        "sektörlerinde ortalama 2,5 ilave iş yarattığını kanıtlamaktadır (UNIDO, 2024). Berkeley Üniversitesi "
        "ekonomisti Enrico Moretti (2010), yüksek teknoloji imalatında bu çarpanın 4,9'a ulaştığını "
        "hesaplamıştır. Bu bulgu, imalat istihdamına yönelik her teşvik TL'sinin ekonomi genelinde çok "
        "katmanlı bir istihdam ve gelir etkisi yaratacağını ortaya koymakta; söz konusu sektörü istihdam "
        "politikasının odak noktası haline getirmektedir."
    )

    add_para(
        "Ancak istihdam teşvikleri tek başına yeterli değildir. Dünya Ekonomik Forumu'nun (WEF) Mesleklerin "
        "Geleceği 2025 Raporu, 2030'a kadar mevcut iş becerilerinin %39'unun geçerliliğini yitireceğini ve "
        "işverenlerin %41'inin yapay zeka otomasyonu nedeniyle işgücü azaltımını planladığını ortaya "
        "koymaktadır (WEF, 2025). Bu tablo, istihdam teşviklerinin aktif işgücü programları —mesleki "
        "yeniden eğitim, İŞKUR sektörel kursları, işbaşı eğitimi, dijital beceri geliştirme— ile "
        "eş zamanlı ve entegre biçimde tasarlanmasını zorunlu kılmaktadır. Card, Kluve ve Weber'in (2018) "
        "200'den fazla ALMP programını kapsayan meta-analizi, mesleki eğitim programlarının orta vadede "
        "anlamlı istihdam ve kazanç artışları sağladığını, iş arama yardımının kısa vadede etkili "
        "olduğunu ve bu iki aracın finansal teşviklerle entegre edildiğinde sinerjik etki yarattığını "
        "belgelemektedir (Card, Kluve ve Weber, 2018)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. GİRİŞ VE AMAÇ
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. GİRİŞ VE AMAÇ", level=1)

    add_para(
        "Küresel ekonomi, COVID-19 sonrası toparlanma sürecinden çıkarak yeni bir döneme girmektedir. "
        "Jeopolitik gerilimler, ticaret savaşları, iklim değişikliği ve teknolojik dönüşüm, ülkelerin "
        "sektörel stratejilerini yeniden şekillendirmesini zorunlu kılmaktadır. Bu süreçte istihdam "
        "politikalarının sektörel bazda tasarlanması, hem sosyal hem ekonomik açıdan kritik bir öneme "
        "sahiptir."
    )

    add_para(
        "Türkiye, OECD'nin en hızlı büyüyen ekonomilerinden biri olmakla birlikte, verimlilik artışının "
        "yavaşlaması, kayıt dışılığın yüksekliği, kadın işgücü katılımının düşüklüğü ve beceri "
        "uyumsuzlukları gibi yapısal sorunlarla karşı karşıyadır (OECD, 2025). İmalat sanayinin orta "
        "teknolojiye yoğunlaşmış yapısı, küresel değer zincirlerinde yukarı doğru entegrasyonu "
        "sınırlandırmaktadır."
    )

    add_para(
        "Bu raporun amacı; (i) TÜİK ve SGK verilerine dayalı olarak sektörlerin ürettiği katma değer "
        "ile işgücü maliyetini karşılaştırmak, (ii) istihdam kapasitesi yüksek ancak işgücü maliyeti "
        "de yüksek olan sektörleri tespit etmek, (iii) uluslararası kurumların sektörel tahminleri ve "
        "küresel eğilimler ışığında desteklenmesi gereken sektörleri belirlemek ve (iv) bir istihdam "
        "teşviki hayata geçirilecekse dikkat edilmesi gereken parametreleri ortaya koymaktır."
    )

    add_para(
        "Rapor, IMF Dünya Ekonomik Görünümü (Ekim 2025, Ocak 2026), OECD Türkiye Ekonomik "
        "İncelemesi (2025), Dünya Bankası Ülke Ekonomik Memorandumu, ILO Dünya İstihdam ve Sosyal "
        "Görünüm raporu (2025), UNIDO Sanayi İstatistikleri Yıllığı (2025) ile TÜİK ve SGK 2024 yılı "
        "verilerini bir arada değerlendirmektedir."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. KÜRESEL EKONOMİK GÖRÜNÜM
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. KÜRESEL EKONOMİK GÖRÜNÜM VE SEKTÖREL TRENDLER", level=1)

    doc.add_heading("3.1. Makroekonomik Çerçeve", level=2)

    add_para(
        "IMF'nin Ocak 2026 güncellemesine göre küresel büyüme 2026'da %3,3, 2027'de %3,2 olarak "
        "öngörülmektedir. Teknoloji yatırımları, mali ve parasal destekler, esnek finansal koşullar "
        "ve özel sektör adaptasyonu, ticaret politikası değişikliklerini dengelemektedir. Küresel "
        "enflasyonun düşmesi beklenmekte, ancak ABD enflasyonunun hedefe daha yavaş dönmesi "
        "öngörülmektedir (IMF, 2026)."
    )

    # IMF Tablo
    add_para("Tablo 1. IMF Küresel Büyüme Tahminleri (%)", bold=True, indent=False)
    t = doc.add_table(rows=7, cols=4)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Bölge/Ülke", "2024", "2025", "2026"]
    imf_data = [
        ["Küresel", "3,3", "3,2", "3,3"],
        ["Gelişmiş Ekonomiler", "—", "1,6", "1,6"],
        ["ABD", "2,8", "2,0", "2,1"],
        ["Euro Bölgesi", "—", "1,2", "1,1"],
        ["Gelişmekte Olan", "4,3", "4,2", "4,0"],
        ["Çin", "—", "4,8", "4,2"],
    ]
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row_data in enumerate(imf_data):
        for j, val in enumerate(row_data):
            t.rows[i+1].cells[j].text = val
    add_para("Kaynak: IMF World Economic Outlook, Ocak 2026; Ekim 2025.", italic=True, indent=False)
    add_para(
        "Not: IMF sınıflandırmasında 'Gelişmiş Ekonomiler' kategorisi; ABD, Almanya, Fransa, "
        "İngiltere, Japonya, Kanada, İtalya, Avustralya, Güney Kore, İspanya, Hollanda, Belçika, "
        "İsviçre, İsveç, Norveç, Danimarka, Avusturya, Finlandiya, Yeni Zelanda, İrlanda ve diğer "
        "yüksek gelirli OECD üyelerini kapsamaktadır. Türkiye IMF sınıflandırmasında 'Gelişmekte "
        "Olan Piyasalar ve Gelişmekte Olan Ekonomiler' grubunda yer almaktadır.",
        italic=True, indent=False
    )

    add_para(
        "Ekim 2025 WEO'sunda \"Değişim İçinde Küresel Ekonomi, Görünüm Karanlık\" başlığıyla yayımlanan "
        "rapor, büyüme tahminlerini 2025 için %3,2, 2026 için %3,1 olarak belirlemiştir. Ocak 2026 "
        "güncellemesiyle 2026 tahmini %3,3'e yukarı yönlü revize edilmiştir. Riskler aşağı yönlü "
        "olmaya devam etmekte; uzayan belirsizlik, artan korumacılık ve işgücü arz şokları büyümeyi "
        "aşağı çekebilir (IMF, 2025a)."
    )

    doc.add_heading("3.2. Sektörel Dönüşüm Dinamikleri", level=2)

    add_para(
        "Küresel düzeyde beş temel sektörel dönüşüm dinamiği ön plana çıkmaktadır:"
    )

    dynamics = [
        ("Yapay Zeka ve Teknoloji Yatırımları: ",
         "IMF, mevcut yapay zeka patlamasını 1990'ların dot-com patlamasına benzetmektedir. "
         "Özellikle Kuzey Amerika ve Asya'da yoğunlaşan AI odaklı yatırımlar büyümeyi ve hisse "
         "senedi piyasalarını desteklemektedir. Ancak verimlilik kazanımları gerçekleşmezse, keskin "
         "piyasa düzeltmeleri ve hanehalkı servetinde azalma riski bulunmaktadır (IMF, 2025a)."),
        ("Yeşil Dönüşüm ve CBAM: ",
         "AB Karbon Sınır Düzenleme Mekanizması (CBAM) 1 Ocak 2026'da yürürlüğe girmiştir. "
         "Çimento, demir-çelik, alüminyum, gübre, elektrik ve hidrojen ithalatına uygulanan mekanizma, "
         "Türkiye'yi en çok etkilenen ülkeler arasına koymaktadır. AB'ye yılda yaklaşık 8 milyar Euro "
         "değerinde karbon yoğun ürün ihraç eden Türkiye, Çin'den sonra ikinci sırada yer almaktadır "
         "(European Commission, 2025; WEF, 2025)."),
        ("Ticaret ve İmalat Yapısı: ",
         "Çin'in ileri imalat, temiz enerji ve ihracat çeşitlendirmesine yönelmesi, emtia ve sermaye "
         "malı talebini dönüştürmekte; teknoloji ve elektrikli araçlarda rekabeti yoğunlaştırmaktadır. "
         "UNIDO verilerine göre küresel imalat 2024'te %2,9 büyümüş, orta-yüksek teknolojili sanayiler "
         "yükselmeye devam etmektedir (UNIDO, 2025)."),
        ("Sanayi Politikası Yeniden Yükselişi: ",
         "Ülkeler giderek daha fazla stratejik sektörleri ve firmaları desteklemek için sanayi politikası "
         "araçlarına başvurmaktadır. Motivasyonlar arasında verimliliği artırma, ithalata bağımlılığı "
         "azaltma ve dayanıklılığı güçlendirme yer almaktadır (IMF, 2025a)."),
        ("İşgücü Piyasası Dönüşümü: ",
         "ILO'ya göre küresel istihdam büyümesi 2025'te yalnızca %1,5 olarak tahmin edilmektedir. "
         "İmalat istihdamında uzun süreli bir durgunluk yaşanmakta, yapısal dönüşüm yüksek katma "
         "değerli hizmetlere doğru yavaşlamaktadır. Beceri uyumsuzlukları verimlilik kazanımlarını "
         "kısıtlamaktadır (ILO, 2025)."),
    ]
    for title, desc in dynamics:
        p = doc.add_paragraph()
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(desc)

    doc.add_heading("3.3. İşgücü Piyasası Küresel Eğilimler", level=2)

    add_para(
        "ILO'nun Dünya İstihdam ve Sosyal Görünüm 2025 raporuna göre, 2024'te küresel istihdam işgücü "
        "ile orantılı büyümüş ve işsizlik oranı %5 düzeyinde sabit kalmıştır. Ancak genç işsizliği "
        "%12,6 ile yüksek seyretmekte, kayıt dışı çalışma ve çalışan yoksulluğu pandemi öncesi "
        "seviyelerine dönmüştür. Düşük gelirli ülkeler en büyük zorluklarla karşı karşıyadır (ILO, 2025)."
    )

    add_para(
        "OECD ülkelerinde işgücü verimliliği büyümesi 2024'te ortalama %0,4 ile zayıf kalmıştır. "
        "Son on yılda verimlilik artışı yavaşlamış, yüksek ve düşük gelirli ülkeler arasındaki "
        "uçurum genişlemiştir. Otomasyon ve dijitalleşme belirli sektörlerde verimliliği artırırken, "
        "benimseme endüstriler ve bölgeler arasında eşitsiz kalmaktadır (OECD, 2025b)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. TÜRKİYE EKONOMİSİ: SEKTÖREL YAPI ANALİZİ
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. TÜRKİYE EKONOMİSİ: SEKTÖREL YAPI ANALİZİ", level=1)

    doc.add_heading("4.1. Gayrisafi Katma Değer ve İşgücü Maliyeti", level=2)

    add_para(
        "TÜİK verilerine göre 2024 yılında gelir yöntemiyle hesaplanan gayrisafi katma değer (cari "
        f"fiyatlarla) {toplam_gkd:,.1f} milyar TL, işgücüne yapılan ödemeler ise "
        f"{summary['isgucu_2024'].sum():,.1f} milyar TL olarak gerçekleşmiştir. İşgücü maliyet "
        f"oranı (işgücü ödemeleri/GKD) ekonomi genelinde %{ort_maliyet:.1f}'dir."
    )

    add_para("Tablo 2. Sektörel Gayrisafi Katma Değer ve İşgücü Maliyeti (2024)", bold=True, indent=False)
    cols = ["Sektör", "GKD\n(Myr TL)", "İşgücü\n(Myr TL)", "Maliyet\nOranı(%)", "İstihdam", "Kadran"]
    keys = ["sektor", "gkd_2024", "isgucu_2024", "maliyet_orani_2024", "istihdam", "kadran"]

    t2 = doc.add_table(rows=1 + len(summary), cols=len(cols))
    t2.style = "Light Grid Accent 1"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(cols):
        cell = t2.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, (_, row) in enumerate(summary.iterrows()):
        for j, key in enumerate(keys):
            val = row.get(key, "")
            if key in ("gkd_2024", "isgucu_2024"):
                text = f"{val:,.1f}"
            elif key == "maliyet_orani_2024":
                text = f"{val:.1f}"
            elif key == "istihdam":
                text = f"{val:,.0f}"
            elif key == "kadran":
                text = str(val)[:2]
            else:
                text = str(val)
            cell = t2.rows[1+i].cells[j]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)

    add_para("Kaynak: TÜİK GSYH Gelir Yöntemi (2024), SGK Yıllık İstatistikleri (2024).", italic=True, indent=False)
    add_para("Not: GKD = Gayrisafi Katma Değer. Maliyet Oranı = İşgücü Ödemesi / GKD × 100.", italic=True, indent=False)

    doc.add_heading("4.2. İstihdam Yapısı ve Verimlilik", level=2)

    add_para(
        f"SGK 4/1-a kapsamında toplam {toplam_istihdam:,.0f} zorunlu sigortalı bulunmaktadır. "
        f"İmalat sanayi %{summary[summary['sektor'].str.contains('malat')]['istihdam_payi'].iloc[0]:.1f} ile "
        f"en büyük istihdam payına sahipken, onu sırasıyla toptan ve perakende ticaret ile inşaat "
        f"izlemektedir. Bu üç sektör toplam istihdamın yaklaşık üçte ikisini oluşturmaktadır."
    )

    add_para(
        "Kişi başı katma değer açısından en verimli sektörler gayrimenkul, ulaştırma ve tarım iken; "
        "kişi başı işgücü maliyeti en yüksek sektörler kamu yönetimi, eğitim ve sağlık olarak "
        "öne çıkmaktadır. Bu durum, emek yoğun kamu hizmet sektörlerinde işgücü maliyetinin katma "
        "değere oranla yüksek olduğuna işaret etmektedir."
    )

    add_para(
        "OECD (2025) raporuna göre Türkiye'de çalışan başına potansiyel büyüme yavaşlamakta ve "
        "görece düşük kalmaktadır. Demografik temettünün sona ermesinin ardından büyümeyi desteklemek "
        "için özellikle hizmetler sektöründe verimlilik artışı, işgücünün nitelik geliştirilmesi, "
        "yenilikçiliğin güçlendirilmesi ve iş düzenlemelerinin kolaylaştırılması gerekmektedir."
    )

    doc.add_heading("4.3. Kadran Analizi: Sektörel Sınıflandırma", level=2)

    add_para(
        "Sektörler, istihdam payı ve işgücü maliyet oranının medyan değerlerine göre dört kadrana "
        "ayrılmıştır. Bu sınıflandırma, teşvik politikalarının hedeflenmesi için analitik bir çerçeve "
        "sunmaktadır:"
    )

    quadrant_descriptions = [
        ("Kadran 1 — Yüksek İstihdam + Yüksek Maliyet (Teşvik Önceliği): ",
         "Bu gruptaki sektörler hem ekonominin istihdam lokomotifi hem de işgücü maliyeti baskısı "
         "altındaki sektörlerdir. İstihdam teşviki bu sektörlerde en yüksek marjinal etkiyi yaratabilir."),
        ("Kadran 2 — Yüksek İstihdam + Düşük Maliyet (Sürdürülebilir Yapı): ",
         "İstihdam kapasitesi yüksek olmasına rağmen işgücü maliyeti görece düşüktür. Mevcut yapının "
         "korunması ve rekabet gücünün desteklenmesi önceliklidir."),
        ("Kadran 3 — Düşük İstihdam + Yüksek Maliyet (Yapısal Dönüşüm): ",
         "İstihdam kapasitesi sınırlı ancak maliyet baskısı yüksektir. Verimlilik artışı, dijital "
         "dönüşüm ve nitelik dönüşümü öncelikli politika alanlarıdır."),
        ("Kadran 4 — Düşük İstihdam + Düşük Maliyet (İzleme): ",
         "Hem istihdam hem maliyet düşüktür. Büyüme potansiyeli izlenmeli, stratejik sektörlerde "
         "hedefli müdahaleler değerlendirilmelidir."),
    ]
    for title, desc in quadrant_descriptions:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(desc)

    # K1 sektörleri tablo
    add_para("Tablo 3. Kadran 1 Sektörleri: Teşvik Öncelikli Sektörler", bold=True, indent=False)
    k1_cols = ["Sektör", "İstihdam\nPayı(%)", "Maliyet\nOranı(%)", "Etkinlik", "KOBİ\n(%)", "Günlük\nKazanç(TL)"]
    k1_keys = ["sektor", "istihdam_payi", "maliyet_orani_2024", "maliyet_etkinligi", "kobi_orani", "ort_gunluk_kazanc"]

    t3 = doc.add_table(rows=1 + len(k1), cols=len(k1_cols))
    t3.style = "Light Grid Accent 1"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(k1_cols):
        t3.rows[0].cells[j].text = h
        for p in t3.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, (_, row) in enumerate(k1.iterrows()):
        for j, key in enumerate(k1_keys):
            val = row.get(key, "")
            if key == "sektor":
                text = str(val)
            elif key == "ort_gunluk_kazanc":
                text = f"{val:,.0f}"
            elif key == "maliyet_etkinligi":
                text = f"{val:.2f}"
            else:
                text = f"{val:.1f}"
            t3.rows[1+i].cells[j].text = text
            for p in t3.rows[1+i].cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)

    add_para("Kaynak: TÜİK (2024), SGK (2024). Yazarın hesaplamaları.", italic=True, indent=False)
    add_para(
        "Not: Etkinlik (Maliyet Etkinliği) = Gayrisafi Katma Değer / İşgücü Ödemesi. Katsayı, 1 TL işgücü "
        "maliyetiyle kaç TL gayrisafi katma değer üretildiğini göstermektedir. Örneğin 2,29 değeri, 1 TL "
        "işgücü ödemesine karşılık 2,29 TL katma değer üretildiği anlamına gelir. Değer ne kadar yüksekse "
        "işgücünün katma değer yaratma kapasitesi o kadar güçlüdür; düşük değer ise işgücü maliyetinin "
        "katma değere oranla yüksek olduğuna ve maliyet baskısına işaret eder.",
        italic=True, indent=False
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. ULUSLARARASI KURUM DEĞERLENDİRMELERİ
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. ULUSLARARASI KURUM DEĞERLENDİRMELERİ", level=1)

    doc.add_heading("5.1. IMF Perspektifi", level=2)
    add_para(
        "IMF, Ekim 2025 Dünya Ekonomik Görünümü'nde yapay zeka yatırımlarının büyümeyi desteklediğini "
        "ancak verimlilik kazanımlarının gerçekleşmemesi halinde keskin piyasa düzeltmeleri riskine dikkat "
        "çekmektedir. Sanayi politikasının yeniden yükselişi, ülkelerin stratejik sektörleri desteklemek "
        "için aktif müdahale araçlarına başvurduğuna işaret etmektedir. Göç politikalarındaki sıkılaşma, "
        "özellikle beceri açığı bulunan sektörlerde firmaların yatırım ve işe alım kararlarını "
        "olumsuz etkileyebilir (IMF, 2025a; IMF, 2026)."
    )

    doc.add_heading("5.2. OECD Önerileri", level=2)
    add_para(
        "OECD'nin Nisan 2025 tarihli Türkiye Ekonomik İncelemesi, birçok sektörel öneri içermektedir. "
        "Buna göre Türkiye ekonomisi orta teknolojili sektörlerde uzmanlaşmış durumdadır ve yüksek "
        "becerili imalat ve hizmetlerde rekabet gücünü artırması gerekmektedir. Küresel değer "
        "zincirlerinde yukarı yönlü entegrasyon, inovasyonun yaygınlaştırılması ve işgücü niteliklerinin "
        "geliştirilmesini gerektirmektedir (OECD, 2025)."
    )
    add_para(
        "OECD özellikle şu alanlara dikkat çekmektedir: (i) Kadın işgücü katılımının artırılması — "
        "OECD ortalamasının önemli ölçüde altında olan kadın işgücü katılımının yükseltilmesi için "
        "okul öncesi eğitim olanaklarının genişletilmesi, çocuk yardımlarının artırılması ve ebeveynler "
        "için efektif vergi oranlarının düşürülmesi; (ii) Hizmetler sektöründe düzenleyici reformlar "
        "(aşağıda ayrıca ele alınmaktadır); (iii) Kayıt dışılığın azaltılması — düşük gelirli "
        "çalışanlarda sosyal güvenlik primlerinin düşürülmesi işgücü piyasası katılımını artırabilir "
        "ve kayıt dışılığı azaltabilir (OECD, 2025)."
    )
    add_para(
        "Hizmetler sektöründe düzenleyici reformlar meselesine ayrıca değinmek gerekmektedir. Türkiye'de "
        "avukatlık, mali müşavirlik, mühendislik, mimarlık gibi mesleki hizmet faaliyetlerine giriş, "
        "OECD ülkeleri arasında en katı lisans, kota ve faaliyet kısıtlamalarına tabidir (OECD Product "
        "Market Regulation Indicators, 2015). Bu durum rekabeti kısıtlamakta, fiyatları yapay olarak "
        "yüksek tutmakta ve verimliliği baskılamaktadır. Ayrıca bu sektörler birçok üretim faaliyetinin "
        "tedarikçisi konumunda olduğundan —hukuki danışmanlık, teknik hizmetler, lojistik destek— "
        "mesleki hizmetlerdeki yüksek maliyet ve düşük verimlilik, değer zincirinin yukarı akışına, "
        "yani imalat ve diğer sektörlere de aktarılmaktadır. Bu nedenle yalnızca hizmetler sektörü "
        "değil, tüm ekonomi verimliliği olumsuz etkilenmektedir."
    )

    doc.add_heading("5.3. Dünya Bankası Değerlendirmesi", level=2)
    add_para(
        "Dünya Bankası'nın \"İstihdam ile Refah\" başlıklı Türkiye Ülke Ekonomik Memorandumu, "
        "Türkiye'nin gelir yakınsama hızının son 15 yılda küresel düzeyde en dikkat çekici örneklerden "
        "biri olduğunu belirtmektedir. Gelir yakınsama (income convergence), düşük gelirli ülkelerin "
        "kişi başına düşen gelirinin zengin ülkelerin gelir düzeyine doğru yaklaşması anlamına gelir. "
        "Dünya Bankası bu bulgusunda Türkiye'nin kişi başı GSYH'sinin ABD ve Almanya gibi gelişmiş "
        "ekonomilerin kişi başı GSYH'sine oranının son 15 yılda tarihsel ortalamaların üzerinde hızla "
        "arttığını vurgulamaktadır. Bu, Türkiye'nin söz konusu süreçte nispeten yoksulluktan nispeten "
        "yüksek gelire doğru küresel ölçekte en başarılı örneklerden biri olduğu anlamına gelir. "
        "Ancak son dönemde bu yakınsamanın hız kestiğine dair işaretler, yüksek enflasyon, düşük "
        "verimlilik artışı ve zayıflayan doğrudan yabancı yatırım gibi uzun süreli yapısal zorluklarla "
        "birlikte potansiyel büyümeyi aşındırmaktadır (World Bank, 2025)."
    )
    add_para(
        "Dünya Bankası, iklime duyarlı tarımın güçlendirilmesi, sanayinin yeşillendirilmesi ve "
        "ihracata yönelik imalatta —otomotiv, makine, metal ve elektronik— istihdam artışının "
        "desteklenmesi gerektiğini vurgulamaktadır. İşsizlik oranı %8 düzeyinde düşük görünmesine "
        "rağmen, geniş tanımlı eksik istihdam oranı (cesareti kırılmış ve eksik çalışanlar dahil) "
        "%29,6 ile oldukça yüksek seyretmektedir (World Bank, 2025)."
    )

    doc.add_heading("5.4. ILO ve UNIDO Bulguları", level=2)
    add_para(
        "ILO'nun 2025 raporuna göre imalat sektörü, bir dönem ekonomik büyüme ve verimli istihdamın "
        "itici gücü iken, uzun süreli bir durgunluk dönemine girmiştir. Bu durum hem ülkeler arasında "
        "hem de ülke içinde mekansal eşitsizlikleri derinleştirmektedir. Buradaki mekansal eşitsizlik, "
        "imalat istihdamının durgunlaşmasının coğrafi boyutunu ifade etmektedir. Ülkeler arasında "
        "bakıldığında, imalat sanayi gücünü kaybeden ülkeler (örneğin geleneksel Batı Avrupa ekonomileri) "
        "ile imalatı büyüyen ekonomiler (Asya ülkeleri) arasındaki gelir ve refah farkı artmaktadır. "
        "Ülke içinde ise sanayi tesislerinin yoğunlaştığı bölgeler ile imalat istihdamının erişilmediği "
        "kırsal veya geri kalmış bölgeler arasındaki ekonomik fark derinleşmektedir. Türkiye için somut "
        "örnek: İstanbul, Bursa, Kocaeli, İzmir gibi sanayi merkezleri ile Doğu ve Güneydoğu Anadolu "
        "illeri arasındaki istihdam ve gelir uçurumu bu eşitsizliğin tipik bir yansımasıdır. Yapısal "
        "dönüşümün imalat ve yüksek katma değerli hizmetlere doğru yavaşlaması, reel harcanabilir "
        "gelirlerin enflasyon dönemindeki kayıplardan toparlanmasını engellemektedir (ILO, 2025)."
    )
    add_para(
        "UNIDO'nun Sanayi İstatistikleri Yıllığı 2025'e göre küresel imalat 2024'te %2,9 büyümüştür. "
        "Asya ve Okyanusya'ya doğru yapısal kayma sürmekte, orta ve yüksek teknolojili sanayiler "
        "yükselmektedir. Güçlü imalat kapasitesine ve çeşitlendirilmiş sanayi yapısına sahip ülkelerin "
        "COVID-19 pandemisinin hem ekonomik hem sağlık etkilerini akranlarından daha iyi atlattığı "
        "belirlenmiştir (UNIDO, 2025)."
    )
    add_para(
        "UNIDO Sanayi Kalkınma Raporu 2024, imalat istihdamının çarpan etkisini kantitatif olarak "
        "belgelemiştir: her imalat istihdamı ekonominin diğer sektörlerinde ortalama 2,5 ilave iş "
        "yaratmaktadır (UNIDO, 2024). Bu sonuç, bağımsız akademik çalışmalarla da teyit edilmektedir. "
        "Berkeley Üniversitesi'nden ekonomist Enrico Moretti, tradable (ihraç edilebilir) sektörlerdeki "
        "her yeni istihdam için yerel hizmet sektöründe 1,6 ilave iş yaratıldığını tespit etmiş; yüksek "
        "teknoloji imalatında bu çarpanın 4,9'a kadar çıktığını ortaya koymuştur (Moretti, 2010). "
        "Brookings Institution ise ileri teknoloji sanayilerindeki her istihdamın ulusal çapta 2,2 "
        "dolaylı iş yarattığını bulmuştur (Brookings, 2015). ABD İmalatçılar Ulusal Birliği (NAM) "
        "hesaplamalarına göre imalat sektöründe harcanan her 1 dolara karşılık ekonomiye 1,89 dolar "
        "katkı sağlanmaktadır (NAM, 2025). Bu bulgular Türkiye için kritik bir politika çıkarımı "
        "sunmaktadır: imalat istihdamına yönelik her teşvik TL'si, yalnızca bu sektörde değil tüm "
        "ekonomide çok katmanlı bir istihdam ve gelir etkisi yaratmaktadır."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. STRATEJİK ÖNCELİKLİ SEKTÖRLER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. STRATEJİK ÖNCELİKLİ SEKTÖRLER", level=1)

    doc.add_heading("6.1. Desteklenmesi Gereken Sektörler ve Gerekçeleri", level=2)

    add_para(
        "Veri analizimiz, uluslararası kurum önerileri ve küresel eğilimler bir arada "
        "değerlendirildiğinde, Türkiye'de öncelikli olarak desteklenmesi gereken sektörler ve "
        "gerekçeleri aşağıda sunulmaktadır:"
    )

    priority_sectors = [
        ("1. İmalat Sanayi",
         "Toplam istihdamın %25'ini oluşturan imalat sanayi, %43,6 işgücü maliyet oranı ile "
         "Kadran 1'de yer almaktadır. UNIDO'ya göre her imalat istihdamı ekonominin diğer "
         "sektörlerinde 2+ iş yaratmaktadır. AB CBAM'ın doğrudan etkileyeceği demir-çelik, "
         "çimento ve alüminyum gibi alt sektörlerde yeşil dönüşüm yatırımları kritiktir. "
         "Türkiye'nin Orta Vadeli Programı (2026-2028) kapsamında HIT-30 Yüksek Teknoloji "
         "Yatırım Programı bu sektörü doğrudan hedeflemektedir. "
         "Teşvik türü: SGK prim indirimi + yeşil dönüşüm yatırım desteği + Ar-Ge istihdamı teşviki."),

        ("2. İnşaat",
         "1,99 milyon sigortalı ile istihdamın %12'sini oluşturmaktadır. İşgücü maliyet oranı "
         "%35,6 ile yüksek, kayıt dışılık sektörün en önemli sorunudur. Kadın istihdam oranı "
         "yalnızca %8,3 ile son derece düşüktür. Deprem sonrası yeniden yapılanma süreci devam "
         "etmekte, yeşil bina standartları ve enerji verimliliği yatırımları önem kazanmaktadır. "
         "Teşvik türü: Kayıt dışından kayıt içine geçiş teşviki + kadın istihdamı ek primi + "
         "yeşil bina sertifikasyonu desteği."),

        ("3. Konaklama ve Yiyecek Hizmeti",
         "1,17 milyon sigortalı ile turizm ekosisteminin temelini oluşturmaktadır. %40,1 maliyet "
         "oranı ve %99,3 KOBİ yoğunluğu ile mikro işletme ağırlıklı bir yapıya sahiptir. "
         "Mevsimsellik ve kayıt dışılık temel sorunlardır. Türkiye'nin turizm gelirlerinin "
         "GSYH'deki payı göz önüne alındığında, bu sektörün desteklenmesi dış denge açısından "
         "da önemlidir. "
         "Teşvik türü: Sezon dışı istihdam desteği + KOBİ dijitalleşme teşviki + "
         "kayıt dışılık denetimi ile eşzamanlı prim indirimi. "
         "KOBİ dijitalleşme teşviki kapsamında; sektördeki mikro ve küçük işletmelerin büyük "
         "çoğunluğu hâlâ nakit esaslı, manuel süreçlerle çalışmaktadır. Bu durum hem kayıt "
         "dışılığa zemin hazırlamakta hem de verimliliği kısıtlamaktadır. Söz konusu teşvik "
         "kapsamında: (i) online rezervasyon ve otomatik ödeme sistemlerinin kurulumu, "
         "(ii) e-fatura ve e-defter uygulamalarına geçiş, (iii) müşteri ilişkileri yönetimi "
         "(CRM) yazılımı, (iv) dijital pazarlama ve online görünürlük yatırımları ile "
         "(v) KOSGEB Dijital Dönüşüm Destek Programı kapsamındaki hibe imkânlarından yararlanma "
         "desteklenmelidir. Bu sayede işletmelerin formel sisteme entegrasyonu kolaylaşacak, "
         "SGK bildirgeleri artacak ve kayıt dışı istihdamın tespiti güçleşmek yerine teşvik "
         "mekanizmasıyla azalacaktır."),

        ("4. İnsan Sağlığı ve Sosyal Hizmet",
         "984 bin sigortalı ile %72 işgücü maliyet oranına sahiptir. Nitelikli işgücü yoğun "
         "yapısı nedeniyle maliyet oranı yapısal olarak yüksektir. Yaşlanan nüfus ve kronik "
         "hastalıklar nedeniyle artan talep, bu sektördeki istihdamın sürdürülebilirliğini "
         "stratejik kılmaktadır. OECD, sağlık harcamalarının verimli kullanımı için dijital "
         "sağlık çözümlerini önermektedir. "
         "Teşvik türü: Nitelikli personel istihdamı desteği + dijital sağlık yatırımı teşviki + "
         "kırsal bölge sağlık istihdamı ek primi. "
         "Dijital sağlık yatırımı teşviki kapsamında; telemedicine (uzaktan tıp) ve yapay zeka "
         "destekli tanı sistemleri yoluyla hem maliyetlerin düşürülmesi hem de erişilebilirliğin "
         "artırılması hedeflenmektedir. Sağlık Bakanlığı MHRS (Merkezi Hekim Randevu Sistemi) "
         "altyapısının genişletilmesi, ÇKYS (Çekirdek Kaynak Yönetim Sistemi) entegrasyonu ve "
         "hastane bilgi yönetim sistemlerine yapılan yatırımlar bu kapsamda değerlendirilebilir. "
         "Kırsal bölge sağlık istihdamı ek primi kapsamında ise Türkiye'de hekim ve sağlık "
         "çalışanı dağılımı bölgeler arasında ciddi dengesizlik göstermektedir. İstanbul ve büyük "
         "şehirlerde yoğunlaşan uzman hekim ile hemşire sayısının aksine, kırsal kesimde ve Doğu "
         "illerinde sağlık çalışanı ciddi ölçüde yetersizdir. Kırsal bölgelerde görev yapan sağlık "
         "personeline uygulanan ek SGK prim indirimi, hem kırsal sağlık erişimini artıracak hem de "
         "bu bölgelerde formel istihdamı teşvik edecektir. Bu yaklaşım Sağlık Bakanlığı'nın Aile "
         "Hekimliği Teşvik Sistemi ile entegre biçimde tasarlanabilir."),

        ("5. Eğitim",
         "654 bin sigortalı ile %75,3 maliyet oranına sahiptir. Beşeri sermaye birikiminin "
         "temel sektörü olarak uzun vadeli büyüme altyapısını oluşturmaktadır. OECD, Türkiye'nin "
         "GSYH'nin yalnızca %0,3'ünü okul öncesi eğitime ayırdığını (OECD ortalaması %0,8) "
         "ve bu alandaki yatırımın kişi başı geliri %6'ya kadar artırabileceğini hesaplamaktadır "
         "(OECD, 2025). "
         "Teşvik türü: Okul öncesi eğitim yatırım desteği + STEM eğitimci istihdamı teşviki."),

        ("6. İdari ve Destek Hizmet Faaliyetleri",
         "1,43 milyon sigortalı ile %60,5 maliyet oranına sahiptir. Güvenlik, temizlik, "
         "çağrı merkezi gibi emek yoğun faaliyetleri kapsamaktadır. Bu sektör, kadın istihdamı "
         "potansiyeli yüksek (%38,3 kadın oranı) ancak ücret düzeyi düşük bir yapıya sahiptir. "
         "Teşvik türü: Düşük ücretli istihdam SGK prim indirimi + nitelik geliştirme programları. "
         "Düşük ücretli istihdam SGK prim indirimi kapsamında; sektördeki istihdamın büyük bölümü "
         "asgari ücret veya yakınında seyreden düşük ücretlidir. OECD (2025), Türkiye'de düşük "
         "gelirli çalışanlar için işveren SGK prim yükünün azaltılmasının hem formel istihdamı "
         "artıracağını hem de kayıt dışılığı düşüreceğini öngörmektedir. Bu nedenle sektörde "
         "asgari ücretin 1,25 katına kadar olan ücret bantlarında işveren SGK payında kademeli "
         "indirim tasarlanması önerilir. Nitelik geliştirme programları kapsamında ise sektör "
         "istihdamının kalıcı ve üretken kılınması için İŞKUR bünyesinde; güvenlik sertifikasyonu, "
         "ileri temizlik teknolojileri, çağrı merkezi operatörlüğü ve müşteri deneyimi yönetimi "
         "gibi alanlarda sertifikalı kısa meslek eğitim programları oluşturulmalı; bu eğitimleri "
         "tamamlayan çalışanların işverenine ek prim teşviki tanınmalıdır. ILO'nun (2025) bulguları, "
         "ALMP programlarıyla desteklenen istihdam teşviklerinin salt finansal teşviklere kıyasla "
         "çok daha kalıcı istihdam etkisi yarattığını ortaya koymaktadır."),

        ("7. Mesleki, Bilimsel ve Teknik Faaliyetler",
         "678 bin sigortalı ile %41,6 maliyet oranı ve en yüksek kadın oranlarından birine "
         "(%45,6) sahiptir. Mühendislik, Ar-Ge, danışmanlık gibi yüksek katma değerli "
         "faaliyetleri kapsamaktadır. Türkiye'nin Ulusal Yapay Zeka Stratejisi bu sektördeki "
         "istihdamı 50.000 kişiye çıkarmayı hedeflemektedir. "
         "Teşvik türü: Ar-Ge istihdamı ek teşviki + start-up ekosistem desteği + "
         "yüksek teknoloji yatırım indirimi. "
         "Ar-Ge istihdamı ek teşviki kapsamında; 5746 sayılı Ar-Ge Kanunu kapsamındaki mevcut "
         "teşviklere ek olarak, Teknoloji Geliştirme Bölgeleri'nde (teknokent) ve Ar-Ge "
         "merkezlerinde istihdam edilen mühendis, doktor ve araştırmacı için işveren SGK payının "
         "%100'e kadar Hazine tarafından karşılanması önerilir. Start-up ekosistem desteği "
         "kapsamında; ilk 3 yılında 10 kişiden az çalışan istihdam eden teknoloji girişimleri için "
         "SGK prim muafiyeti veya 5 yıllık kademeli prim desteği uygulanabilir. Yüksek teknoloji "
         "yatırım indirimi kapsamında ise HIT-30 programı çerçevesinde yarı iletkenler, elektrikli "
         "araç bileşenleri, yeşil enerji teknolojileri ve biyoteknoloji alanlarına yapılan "
         "yatırımlarda istihdam edilen nitelikli personel için yatırım indirimine ek SGK teşviki "
         "önerilmektedir."),

        ("8. Bilgi ve İletişim (Stratejik Büyüme Sektörü)",
         "Kadran 4'te yer almasına rağmen (%1,7 istihdam payı), dijital dönüşümün temel "
         "sektörü olarak stratejik öneme sahiptir. Ortalama günlük kazanç 2.371 TL ile en "
         "yüksek ücret düzeyine sahip sektörlerden biridir. Türkiye'nin 5G altyapısı, "
         "yazılım ihracatı ve yapay zeka hedefleri doğrultusunda bu sektörün desteklenmesi "
         "uzun vadeli dönüştürücü etkiye sahiptir. "
         "Teşvik türü: Yazılım istihdamı SGK muafiyeti + teknokent yatırım desteği + "
         "dijital beceri geliştirme programları. "
         "Yazılımcı istihdamı SGK muafiyeti kapsamında; 4691 sayılı Teknoloji Geliştirme Bölgeleri "
         "Kanunu kapsamındaki mevcut muafiyetin teknokent dışında faaliyet gösteren yazılım odaklı "
         "KOBİ'lere de genişletilmesi, özellikle yazılım ihracatı yapan firmalar için SGK prim "
         "desteği olarak tasarlanabilir. Teknokent yatırım desteği kapsamında mevcut 90+ teknokent "
         "ve 50+ teknopark altyapısının kapasite ve kalite artırımı için altyapı yatırımlarına "
         "destek sağlanması önerilir; özellikle Anadolu'da yeni teknokentlerin kurulması, bölgesel "
         "dijital istihdam dengesizliğinin giderilmesine katkı sağlayacaktır. Dijital beceri "
         "geliştirme programları kapsamında; WEF Future of Jobs 2025 raporu, mevcut becerilerin "
         "%39'unun 2030'a kadar geçerliliğini yitireceğini öngörmektedir. İŞKUR ile yazılım "
         "eğitim sektörü arasında köprü kurularak; yazılım geliştirme, veri bilimi, siber güvenlik "
         "ve bulut altyapısı alanlarında kısa süreli yoğun eğitim programları desteklenmeli ve "
         "mezunların istihdamı SGK teşvikiyle ilişkilendirilmelidir."),
    ]

    for title, desc in priority_sectors:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    doc.add_heading("6.2. Sektörel Risk ve Fırsat Matrisi", level=2)

    add_para("Tablo 4. Desteklenmesi Gereken Sektörlerde Risk ve Fırsat Değerlendirmesi", bold=True, indent=False)
    risk_headers = ["Sektör", "Fırsat", "Risk", "Küresel Trend"]
    risk_data = [
        ["İmalat Sanayi", "CBAM uyumu ile AB pazarı\navantajı; UNIDO çarpan etkisi", "Karbon maliyeti artışı;\nrekabet baskısı", "Yeşil dönüşüm;\norta-yüksek teknoloji"],
        ["İnşaat", "Deprem yeniden yapılanma;\nyeşil bina potansiyeli", "Kayıt dışılık;\nmevsimsellik", "Sürdürülebilir\nyapı standartları"],
        ["Konaklama", "Turizm geliri;\ndış denge katkısı", "Mevsimsellik;\ndüşük ücret", "Dijital turizm;\nsürdürülebilir turizm"],
        ["Sağlık", "Yaşlanan nüfus;\nartan talep", "Nitelikli personel göçü;\nyüksek maliyet", "Dijital sağlık;\ntele-tıp"],
        ["Eğitim", "Beşeri sermaye;\nOECD önerisi", "Bütçe kısıtı;\nnitelik sorunu", "STEM eğitimi;\nhayat boyu öğrenme"],
        ["Bilgi ve İletişim", "AI stratejisi;\nyazılım ihracatı", "Beyin göçü;\nbeceri açığı", "AI patlaması;\n5G altyapı"],
    ]

    t4 = doc.add_table(rows=1 + len(risk_data), cols=len(risk_headers))
    t4.style = "Light Grid Accent 1"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(risk_headers):
        t4.rows[0].cells[j].text = h
        for p in t4.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, rd in enumerate(risk_data):
        for j, val in enumerate(rd):
            t4.rows[i+1].cells[j].text = val
            for p in t4.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    add_para("Kaynak: Yazarın IMF (2025a; 2026), OECD (2025), World Bank (2025), ILO (2025), UNIDO (2024; 2025) verilerine dayalı değerlendirmesi.", italic=True, indent=False)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. İSTİHDAM TEŞVİKİ POLİTİKA ÖNERİLERİ
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. İSTİHDAM TEŞVİKİ POLİTİKA ÖNERİLERİ", level=1)

    doc.add_heading("7.1. Uluslararası Deneyimler ve Vergi Takozu Karşılaştırması", level=2)

    add_para(
        "Vergi takozu nedir? Bir işvereni düşünün: bir çalışanı için ayda 100 TL ödemekte. Ancak "
        "çalışan eline yalnızca 61 TL geçirmekte. Aradaki 39 TL'nin bir kısmı gelir vergisi, bir "
        "kısmı çalışan SGK payı, bir kısmı ise işveren SGK payından oluşmaktadır. İşte bu "
        "\"ödenen ile alınan arasındaki fark\"a vergi takozu denir. Vergi takozunun yüksekliği; "
        "işvereni daha az istihdam yaratmaya, çalışanı ise formel iş yerine kayıt dışı çalışmaya "
        "iter. OECD Taxing Wages 2025 verileri, Türkiye'nin vergi takozunun OECD ortalamasının "
        "belirgin biçimde üzerinde olduğunu ve bu durumun özellikle düşük ücretli işlerde formel "
        "istihdamı caydırıcı bir etken oluşturduğunu ortaya koymaktadır."
    )

    add_para(
        "İstihdam teşviki tasarımında uluslararası deneyimler ve Türkiye'nin vergi takozu konumu "
        "belirleyici bir çerçeve sunmaktadır. OECD Taxing Wages 2025 raporuna göre, Türkiye'de "
        "tek çalışanlı ve ortalama ücretli bir birey için vergi takozu %39,0 olup OECD ortalaması "
        "olan %34,9'un 4,1 puan üzerindedir. İki çocuklu evli aileler için bu fark daha da "
        "dramatiktir: Türkiye %39,0 ile OECD ortalaması %25,7'nin 13,3 puan üzerindedir. Türkiye, "
        "Avrupa'da aileleri bekarlardan farklı vergilendirmeyen tek ülke konumundadır (OECD, 2025d)."
    )

    add_para(
        "Vergi takozunun bileşenlerini anlamak için önce Türkiye'nin gelir vergisi yapısını netleştirmek "
        "gerekir. Türkiye'de 2022 yılından itibaren asgari ücret tutarına karşılık gelen gelir gelir "
        "vergisinden muaf tutulmuştur (Gelir Vergisi Kanunu md.23/18). Dolayısıyla asgari ücret düzeyinde "
        "çalışan kişiler gelir vergisi ödememektedir. Ancak ortalama ücret alanlar —ki bu tutar asgari "
        "ücretin üzerindedir— artan oranlı dilim yapısı (2024: %15, %20, %27, %35, %40) kapsamında "
        "gelir vergisi ödemektedir. OECD Taxing Wages raporunda kullanılan 'ortalama ücretli çalışan' "
        "(%100 APW) profili, bu muafiyetin üzerinde bir gelir düzeyine karşılık geldiğinden Türkiye'nin "
        "%39,0 vergi takozu rakamı gelir vergisini de içermektedir."
    )
    add_para(
        "Vergi takozu hesaplama formülü: Vergi Takozu (%) = (Gelir Vergisi + Çalışan SGK Payı + "
        "İşveren SGK Payı − Nakit Transferler) / Toplam İşgücü Maliyeti × 100. Türkiye'de bileşenler "
        "yaklaşık olarak: Gelir vergisi ~%8-10 (ortalama ücretli için); Çalışan SGK payı %15 "
        "(emeklilik %9 + sağlık %5 + işsizlik %1); İşveren SGK payı %22,75 (fiili uygulanan "
        "teşviklerle ~%17,5'e kadar düşebilir); Nakit transferler sınırlı. Toplam işgücü maliyeti = "
        "brüt ücret + işveren SGK payı. Bu formülün sonucu ortalama ücretli için yaklaşık %39'a "
        "ulaşmaktadır.",
        italic=True, indent=False
    )
    add_para("Tablo 6. OECD Vergi Takozu ve SGK Yükü Karşılaştırması (2024)", bold=True, indent=False)
    tw_table = doc.add_table(rows=7, cols=4)
    tw_table.style = "Light Grid Accent 1"
    tw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tw_headers = ["Gösterge", "Türkiye", "OECD Ortalaması", "Örnekler"]
    tw_data = [
        ["Vergi takozu — bekar, ort. ücret\n(gelir vergisi + çalışan SGK + işveren SGK)",
         "%39,0", "%34,9",
         "Belçika %52,6; Almanya %47,9; Fransa %47,2; İngiltere %30,1"],
        ["Vergi takozu — evli, 2 çocuk, ort. ücret",
         "%39,0", "%25,7",
         "Türkiye aile indirimi tanımayan OECD'de tek ülke"],
        ["  ↳ Gelir vergisi bileşeni (ort. ücret)",
         "~%8-10", "~%13-15",
         "ASGARİ ücret için 0 (2022'den itibaren muafiyet)"],
        ["  ↳ Çalışan SGK payı",
         "%15,0", "~%9-12",
         "Emeklilik %9 + Sağlık %5 + İşsizlik %1"],
        ["  ↳ İşveren SGK payı (yasal)",
         "%22,75", "~%17-20",
         "Fiili (teşvikli): ~%17,5; Fransa ~%28-30"],
        ["Vergi takozunda OECD en yüksek: Belçika",
         "—", "%52,6",
         "Türkiye OECD içinde üst çeyrekte"],
    ]
    for j, h in enumerate(tw_headers):
        tw_table.rows[0].cells[j].text = h
        for p in tw_table.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, rd in enumerate(tw_data):
        for j, val in enumerate(rd):
            tw_table.rows[i+1].cells[j].text = val
            for p in tw_table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    add_para("Kaynak: OECD Taxing Wages 2025; GİB; 193 sayılı GVK md.23/18 (2022 değişikliği).",
             italic=True, indent=False)
    add_para(
        "Not: Türkiye'de vergi takozu yüksekliğinin temel nedeni aile indirimi bulunmaması ve "
        "çalışan SGK payının OECD ortalamasının üzerinde seyretmesidir. Asgari ücret düzeyindeki "
        "çalışanlar gelir vergisinden muaf olduğundan bu grubun vergi takozu daha düşüktür.",
        italic=True, indent=False
    )

    add_para(
        "Uluslararası istihdam teşviki ve istihdam koruma deneyimleri, her ülkenin kendi kurumsal "
        "yapısına özgü tasarımlar geliştirdiğini göstermektedir. Aşağıda dört ülkenin deneyimi "
        "Türkiye ile karşılaştırmalı olarak değerlendirilmektedir."
    )

    intl_experiences = [
        ("Almanya — Kurzarbeit ve Türkiye Karşılaştırması: ",
         "Almanya'nın kısa çalışma programı (Kurzarbeit), 2008-09 küresel krizinde ~500.000, "
         "COVID-19 döneminde 6 milyon işçiyi kapsayarak istihdamı korumuştur. Almanya'da istihdam, "
         "GSYİH'deki %7'lik düşüşe rağmen yalnızca %1 gerilemiştir (Tax Foundation, 2025). "
         "Türkiye'de muadil araç Kısa Çalışma Ödeneği (KÇÖ) olmakla birlikte kapsam ve esneklik "
         "açısından önemli sınırlılıklar taşımaktadır. Buna ek olarak, Türkiye COVID-19 sürecinde "
         "Kısa Çalışma Ödeneği kapsamına giremeyen —özellikle prim koşulunu sağlayamayan— çalışanlar "
         "için Nakdi Ücret Desteği (NÜD) uygulaması başlatmıştır. NÜD, günlük 34,72 TL (dönemin "
         "asgari ücretinin %47'si) tutarında, son 60 günde en az 30 gün sigortalı olan tüm "
         "çalışanlara açık olarak tasarlanmıştır. Bu uygulama KÇÖ'nün kapsam dışında kalan "
         "katmanlara ulaşması bakımından anlamlı ancak ödeme düzeyi yetersiz kalmıştır."),
        ("Fransa — Réduction Générale ve Türkiye Karşılaştırması: ",
         "Yapısal işveren sosyal katkı payı indirimi, 2024-25'te asgari ücretin 1,6 katına kadar "
         "maaşlarda maksimum %32 oranında uygulanmıştır. 2026'da RGDU sistemine dönüştürülmüştür. "
         "Ancak CICE vergi kredisinin ihracat üzerinde istatistiksel olarak anlamlı bir etkisi "
         "bulunamamıştır (France Stratégie, 2020). Türkiye'de 5510/7252 sayılı Kanunlar kapsamındaki "
         "prim indirimleri Fransa'ya kıyasla daha düşük oran ve daha dar kapsama sahiptir; özellikle "
         "aile indirimi boyutu tamamen yoktur."),
        ("Güney Kore — KOBİ İstihdam Teşvikleri ve Türkiye Karşılaştırması: ",
         "Her ek KOBİ çalışanı için 7-7,7 milyon KRW vergi indirimi, genç KOBİ çalışanlarına "
         "5 yıl boyunca %90 gelir vergisi muafiyeti uygulanmaktadır. Ancak yüksek nitelikli "
         "gençlerde mütevazı teşviklerin yüksek rezervasyon ücreti nedeniyle etkisiz kaldığı "
         "gösterilmiştir (MOEL, 2024). Türkiye'de 687 sayılı KHK kapsamında net istihdam artışı "
         "başına SGK teşviki uygulanmaktadır; ancak bu sistem KOBİ'ye özgü değildir ve Güney "
         "Kore'nin vergi indirimi mekanizması gibi sektörel farklılaştırma içermemektedir."),
        ("İngiltere — Employment Allowance ve Türkiye Karşılaştırması: ",
         "Sabit işveren NIC indirimi 2025'te 10.500 GBP'ye yükseltilmiştir. İşveren NIC oranı "
         "%13,8'den %15'e artırılırken tüm işletmelere erişim sağlanmıştır (GOV.UK, 2025). "
         "İngiltere'nin sabit tutarlı (oran değil, mutlak tutar) yapısı, küçük işletmeleri "
         "büyük işletmelere kıyasla orantısız biçimde desteklemektedir. Türkiye'de böyle bir "
         "sabit tutar mekanizması henüz mevcut değildir."),
    ]
    for title, desc in intl_experiences:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(desc)

    add_para("Tablo 7. Almanya Kurzarbeit — Türkiye Kısa Çalışma Ödeneği ve Nakdi Ücret Desteği Karşılaştırması",
             bold=True, indent=False)
    kk_table = doc.add_table(rows=8, cols=3)
    kk_table.style = "Light Grid Accent 1"
    kk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kk_headers = ["Boyut", "Almanya Kurzarbeit", "Türkiye (KÇÖ + NÜD)"]
    kk_data = [
        ["Azami süre", "21 aya kadar (kriz döneminde)", "KÇÖ: 3 ay (uzatılabilir 6 ay)"],
        ["Ödeme oranı", "%60-%80 (kademeli)", "KÇÖ: %60 sabit; NÜD: asgari ücretin ~%47'si"],
        ["Yararlanma eşiği", "İşyerinde çalışanların %10'u etkilenmiş olması",
         "KÇÖ: Son 3 yılda 600 gün prim; NÜD: Son 60 günde 30 gün sigortalı"],
        ["COVID kapsam", "6 milyon çalışan; istihdam yalnızca %1 geriledi",
         "KÇÖ sınırlı kaldı; NÜD ek katman sağladı ancak ödeme düzeyi düşük kaldı"],
        ["Kapsam genişliği", "Tüm SSK katkı ödeyenler",
         "KÇÖ: İşsizlik sigortası primini dolduranlar; NÜD daha geniş ama kısa süreli"],
        ["İşveren katkısı", "İşveren SGK primlerini ödemeye devam eder",
         "İşveren SGK yükümlülüğü devam eder; ek destek sınırlı"],
        ["Değerlendirme", "Güçlü istihdam koruma; esneklik ve kapsam avantajı",
         "KÇÖ+NÜD birlikte anlamlı ancak kapsam ve ödeme yetersizliği reform ihtiyacı doğuruyor"],
    ]
    for j, h in enumerate(kk_headers):
        kk_table.rows[0].cells[j].text = h
        for p in kk_table.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, rd in enumerate(kk_data):
        for j, val in enumerate(rd):
            kk_table.rows[i+1].cells[j].text = val
            for p in kk_table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    add_para(
        "Politika çıkarımı: Türkiye'nin KÇÖ + NÜD ikili yapısını tek, daha kapsayıcı ve daha yüksek "
        "ödeme oranlı bir kısa çalışma mekanizmasına dönüştürmesi, gelecekteki kriz dönemlerinde "
        "istihdam kaybını önemli ölçüde sınırlayabilir. Bu reformda Kurzarbeit'ın işyeri eşiği "
        "(çalışanların %10'u) ve kademeli ödeme (süreye bağlı %60-80) modeli referans alınabilir.",
        italic=True, indent=False
    )

    add_para(
        "Ölü ağırlık kaybı (deadweight loss) nedir? İstihdam teşviklerinde ölü ağırlık kaybı, "
        "\"zaten olacak işler için para harcamak\" anlamına gelir. Örneğin bir firma, teşvik "
        "programı olsa da olmasa da bir mühendisi işe alacaktı. Teşvik verildiğinde devlet para "
        "harcamış, ancak net yeni istihdam yaratılmamıştır. Bu boşa harcanan pay, ölü ağırlık "
        "kaybıdır."
    )
    add_para(
        "Türkiye'ye özgü ölü ağırlık kaybı analizi: Uysal (2016), IZA Discussion Paper No. 9993 "
        "kapsamında Türkiye'nin bölgesel SGK prim teşviklerini fark-içi-farklar (difference-in-"
        "differences) yöntemiyle değerlendirmiştir. Çalışmada iki farklı teşvik programı "
        "incelenmiştir: (1) 5084 sayılı Kanun kapsamındaki ilk program (2004-2012): 49 geri kalmış "
        "ilde uygulanan, işveren SGK payının tamamını kapsayan yoğun destek. Ölü ağırlık kaybı "
        "%47-78 aralığında ölçülmüştür — yani her 100 TL teşvikin 47-78 TL'si zaten gerçekleşecek "
        "istihdama gitmiştir. (2) 5510 sayılı Kanun kapsamındaki ikinci program (2008 sonrası): "
        "Net istihdam artışı koşuluna bağlanan, daha iyi tasarlanmış yapı. Bu programda ölü ağırlık "
        "kaybı %23-44'e gerilemiştir. Temel bulgu: Net istihdam artışı koşulu, ölü ağırlık kaybını "
        "yaklaşık yarıya indirmiştir. Bu sonuç, teşvik programı tasarımında en kritik değişkenin "
        "'net artış şartı' olduğunu kanıtlamakta ve Türkiye için tasarlanacak tüm yeni programlara "
        "doğrudan politika çıkarımı sunmaktadır (Uysal, 2016; IZA DP No. 9993)."
    )

    add_para(
        "Ölü ağırlık kaybı (deadweight loss) tahminleri ülkeden ülkeye önemli farklılıklar "
        "göstermektedir. Türkiye için %27-46, Belçika için %53, Fransa için %84 oranında ölü "
        "ağırlık kaybı hesaplanmıştır. Makroekonomik çalışmalar sistematik olarak anket "
        "çalışmalarından daha yüksek ölü ağırlık kaybı tahmin etmektedir. Teşviklerin yalnızca "
        "işsiz bireylere yönlendirilmesi (işten işe geçişlerin hariç tutulması), Almanya "
        "deneyiminde maliyeti %20-30 düşürmüştür (Hartmann ve diğerleri, 2013; IZA, 2023)."
    )

    add_para(
        "Meksika'da uygulanan deneysel geçici ücret teşvikleri, formel istihdamı 4,2 yüzde puanı "
        "(%14,5) ve daimi sözleşmeleri %25 oranında artırmıştır. Bu bulgu, kayıt dışılığın yüksek "
        "olduğu Türkiye için kayıt dışından kayıt içine geçiş odaklı teşvik tasarımının etkinliğine "
        "güçlü bir kanıt sunmaktadır (VoxDev, 2024; OECD, 2024b)."
    )

    doc.add_heading("7.2. Teşvik Tasarım Parametreleri", level=2)

    add_para(
        "Uluslararası deneyimler ve Türkiye'nin yapısal özellikleri dikkate alındığında, istihdam "
        "teşviki tasarımında aşağıdaki parametrelerin gözetilmesi önerilmektedir:"
    )

    params = [
        ("SGK Prim İndirimi Kademelendirmesi: ",
         "Kadran 1 sektörleri için işveren SGK prim payında %10-15 oranında indirim önerilmektedir. "
         "Net istihdam artışı şartı konulmalı ve baz yıl istihdamının altına düşmeme koşulu "
         "aranmalıdır. OECD Türkiye Ekonomik İncelemesi (2025), sosyal güvenlik katkılarının sabit "
         "oranlı yapısının düşük gelirli çalışanları orantısız biçimde etkilediğini saptamaktadır. "
         "OECD'nin önerisi, asgari ücretin 1,25-1,5 katına kadar ücret bandındaki çalışanlar için "
         "işveren SSK payını kademeli olarak azaltmaktır. Bu sayede işverenlerin düşük ücretli "
         "formel istihdam yaratma maliyeti düşer; kayıt dışı çalışan düşük ücretliler formel sisteme "
         "geçiş için teşvik alır ve asgari ücret etrafındaki formelleşme eşiği problemi azalır. "
         "Fransa bu yaklaşımın başarılı uygulamasına örnek teşkil etmektedir: asgari ücretin 1,6 "
         "katına kadar olan maaşlarda işveren SSK payı maksimum %32 oranında indirime tabi "
         "tutulmakta; bu sayede düşük ücretli sektörlerde istihdam teşvik edilmektedir "
         "(URSSAF, 2025)."),
        ("İşyeri Büyüklüğüne Göre Farklılaştırma: ",
         "İşyeri büyüklüğüne göre farklılaştırmanın gerekçesi şudur: büyük işletmeler, teşvik olsa "
         "da olmasa da yeni çalışan alabilecek mali güce sahiptir; bu nedenle onlara tam oran teşvik "
         "vermek ölü ağırlık kaybını artırır. Küçük ve mikro işletmeler ise işgücü maliyetine en "
         "duyarlı ve kayıt dışılık riski en yüksek gruptur. Önerilen kademeler: Mikro işletmeler "
         "(<10 çalışan) için tam oran teşvik — nakit akışı kısıtlı, istihdam maliyetine en duyarlı "
         "ve kayıt dışılıktan formal geçiş potansiyeli en yüksek grup. Küçük işletmeler (10-49 "
         "çalışan) için %75 teşvik — KOBİ'ler toplam istihdamın büyük bölümünü oluşturur. Orta "
         "işletmeler (50-249 çalışan) için %50 teşvik — maliyet avantajı verimliliğe yatırıma "
         "yönlendirilmeli. Büyük işletmeler (250+ çalışan) için yalnızca net istihdam artışı "
         "koşuluyla teşvik — ölü ağırlık kaybı riski en yüksek grupta yalnızca gerçek yeni "
         "istihdamın teşvik alması sağlanmalıdır. Verilerimize göre K1 sektörlerinde KOBİ oranı "
         "%95'in üzerindedir."),
        ("Bölgesel Farklılaştırma: ",
         "Yatırım teşvik bölgelerine göre ek puan uygulanmalıdır. Dünya Bankası'nın deprem "
         "bölgelerinde ekonomik toparlanma vurgusu (World Bank, 2025) doğrultusunda, deprem "
         "bölgelerinde ek teşvik katsayısı uygulanmalıdır."),
        ("Cinsiyet Eşitliği Boyutu: ",
         "OECD'nin vurguladığı kadın işgücü katılımı açığını kapatmak için kadın oranı %30'un "
         "altındaki sektörlerde (İnşaat %8,3, Ulaştırma, Madencilik) kadın istihdamına ek %5 puan "
         "teşvik uygulanmalıdır. Okul öncesi eğitim yatırımları ile desteklenmelidir (OECD, 2025)."),
        ("Nitelik Düzeyine Göre Ayrım: ",
         "Düşük nitelikli istihdam için kayıt dışılık azaltma odaklı tam teşvik; orta nitelikli "
         "istihdam için standart teşvik; yüksek nitelikli istihdam (mühendis, Ar-Ge uzmanı) için "
         "Türkiye Ulusal AI Stratejisi hedefleri doğrultusunda ek Ar-Ge istihdamı teşviki "
         "uygulanmalıdır."),
        ("Teşvik Süresi ve Kademeli Çıkış: ",
         "Teşvik programlarının açık uçlu uygulanması hem mali sürdürülebilirliği tehdit eder hem "
         "de firmaların teşvike bağımlı hale gelmesine yol açar. Bu nedenle kademeli çıkış "
         "mekanizması baştan tasarlanmalıdır. 1-3. yıl (Tam oran): İşverenin yeni çalışanı işine "
         "alıştırma ve verimli hale getirme sürecini kapsar; tam teşvik ilk istihdam maliyetini "
         "dengeler. 4. yıl (%75'e düşüş): İstihdam ilişkisi yerleşmiş, verimlilik artışı başlamıştır; "
         "%25'lik azalma işverene kademeli uyum süresi tanır. 5. yıl (%50'ye düşüş): Teşvik "
         "bitmeden önce işveren işgücü maliyetini kendi verimliliği ile karşılamaya alışır. "
         "6. yıl (Etki değerlendirmesi): Ulusal istihdam verisi ve teşvik maliyet-fayda analizi "
         "temelinde; programın uzatılması, daraltılması veya sonlandırılması kararı verilir. "
         "OECD iyi uygulama kılavuzu, bu değerlendirmenin bağımsız araştırma kurumu tarafından "
         "yapılmasını önermektedir. Bu tasarım, Almanya'nın Kurzarbeit ve Fransa'nın RGDU "
         "sistemlerinde de uygulanan 'zaman sınırlı teşvik + değerlendirmeye bağlı uzatma' "
         "prensibini izlemektedir."),
    ]
    for title, desc in params:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(desc)

    doc.add_heading("7.3. Sektör Bazlı Teşvik Modelleri", level=2)

    add_para("Tablo 7. Sektör Bazlı Teşvik Model Önerileri", bold=True, indent=False)
    model_headers = ["Sektör", "Teşvik Modeli", "Hedef", "İzleme Göstergesi"]
    model_data = [
        ["İmalat Sanayi", "SGK %12 + yeşil yatırım\ndesteği + Ar-Ge teşviki", "Net 200K yeni istihdam\nCBAM uyumu", "Karbon yoğunluğu\nİhracat payı"],
        ["İnşaat", "Kayıt içi geçiş primi\n+ kadın ek teşvik %5", "Kayıt dışılık %10↓\nKadın oranı %15'e", "SGK bildirge artışı\nKadın sigortalı sayısı"],
        ["Konaklama", "Sezon dışı prim indirimi\n+ dijital KOBİ desteği", "12 ay istihdam sürekliliği\nDijitalleşme %50", "Mevsimsel istihdam farkı\nE-fatura kullanımı"],
        ["Sağlık", "Nitelikli personel primi\n+ kırsal ek teşvik", "Beyin göçünü yavaşlatma\nKırsal erişim", "Hekim sayısı/100K\nHemşire tutma oranı"],
        ["Eğitim", "Okul öncesi yatırım\n+ STEM eğitimci teşviki", "Okul öncesi okullaşma\n%80'e çıkarma", "ECEC harcama/GSYH\nSTEM mezun sayısı"],
        ["Bilgi İletişim", "Yazılımcı SGK muafiyeti\n+ teknokent genişleme", "50K AI istihdamı\nYazılım ihracatı 2x", "Ar-Ge harcama/GSYH\nPatent başvurusu"],
    ]

    t5 = doc.add_table(rows=1 + len(model_data), cols=len(model_headers))
    t5.style = "Light Grid Accent 1"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(model_headers):
        t5.rows[0].cells[j].text = h
        for p in t5.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, rd in enumerate(model_data):
        for j, val in enumerate(rd):
            t5.rows[i+1].cells[j].text = val
            for p in t5.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_heading("7.4. AB Yeşil Mutabakat ve CBAM Uyumu", level=2)

    add_para(
        "AB Karbon Sınır Düzenleme Mekanizması (CBAM) 1 Ocak 2026'da kesin uygulama aşamasına "
        "geçmiştir. Çimento, demir-çelik, alüminyum, gübre, elektrik ve hidrojeni kapsayan mekanizma, "
        "ithalat edilen malların karbon yoğunluğuna dayalı sertifika alınmasını zorunlu kılmaktadır "
        "(European Commission, 2025)."
    )
    add_para(
        "Türkiye'nin CBAM'a maruziyeti somut rakamlarla çarpıcı boyutlara ulaşmaktadır. Uluslararası "
        "Karbon Eylem Ortaklığı (ICAP) ve AB Komisyonu analizlerine göre Türkiye'nin CBAM kapsamındaki "
        "sektörlerdeki AB ihracatı 2022'de 19 milyar Euro değerindedir. CBAM karbon sertifika "
        "maliyetleri 2027'de yaklaşık 138 milyon Euro'ya, karbonun AB ETS fiyatına göre 2032'de "
        "2,5 milyar Euro'ya ulaşabilecektir (ICAP, 2025). Bu maliyet özellikle demir-çelik, çimento "
        "ve alüminyum ihracatçılarını doğrudan vuracaktır. Avrupa Komisyonu analizine göre Türkiye, "
        "Çin'den sonra CBAM'dan en çok etkilenecek ikinci ülkedir (WEF, 2025; IISD, 2025)."
    )
    add_para(
        "Türkiye bu zorluğa iki koldan yanıt vermektedir. Birincisi, 2025 yılında çıkarılan 7552 "
        "sayılı İklim Kanunu ile ulusal ETS'nin yasal çerçevesi oluşturulmuş; yılda 50.000 ton CO₂ "
        "üzerinde emisyon yapan tesisleri kapsayan pilot ETS aşaması 2026-2027 döneminde başlayacaktır. "
        "İkincisi, Yeşil Dönüşüm Destek Programı ve Sanayi ve Teknoloji Bakanlığı'nın yeşil dönüşüm "
        "portalı aracılığıyla düşük karbonlu üretim teknolojileri için Dünya Bankası finansmanlı "
        "1 milyar doların üzerinde destek mobilize edilmektedir. Bu bağlamda iki politika aracının "
        "eş zamanlı uygulanması önerilmektedir: (i) ETS kapsamındaki firmalara CBAM sertifika yüküne "
        "karşılık geçici SGK prim desteği (yeşil geçiş desteği); (ii) düşük karbonlu üretim "
        "teknolojisine yatırım yapan firmalarda istihdam edilen nitelikli personel için ek teşvik. "
        "Aksi takdirde CBAM maliyet yükünün firmaların istihdam azaltımına yol açması riski ciddidir."
    )

    doc.add_heading("7.5. Dijital Dönüşüm ve Yapay Zeka", level=2)

    add_para(
        "Türkiye'nin Ulusal Yapay Zeka Stratejisi (2021-2025, 2024 güncellemesi) yapay zekanın "
        "GSYH'ye katkısını %5'e çıkarmayı ve AI alanında 50.000 istihdam yaratmayı hedeflemektedir. "
        "Orta Vadeli Program (2026-2028), sektörel önceliklendirmeye dayalı hedefli sanayi "
        "politikalarının uygulanacağını, yüksek katma değerli üretim ve teknoloji odaklı dönüşüm "
        "yoluyla rekabet gücünün artırılacağını belirtmektedir."
    )
    add_para(
        "HIT-30 Yüksek Teknoloji Yatırım Programı: Sanayi ve Teknoloji Bakanlığı tarafından Temmuz "
        "2024'te açıklanan program, 2030'a kadar 30 milyar dolarlık devlet desteği ve 20 milyar "
        "dolarlık özel sektör yatırımı hedeflemektedir. Sekiz öncelikli teknoloji alanı "
        "belirlenmiştir: (1) Yarı İletkenler; (2) Elektrikli Araçlar ve Bataryalar; (3) İleri "
        "İmalat ve Robotik; (4) Yeşil Enerji Teknolojileri; (5) Biyoteknoloji ve Sağlık; "
        "(6) Yapay Zeka ve Büyük Dil Modelleri; (7) İletişim Teknolojileri; (8) Uzay "
        "Teknolojileri. Program; vergi teşviklerinin %60'a ulaşabileceği proje bazlı destekler, "
        "Ar-Ge merkezi kuran dünyanın ilk 1000 Ar-Ge şirketi için personel giderlerinin %50'sinin "
        "Hazine tarafından karşılanması ve YTAK aracılığıyla uygun faizli finansman imkânları "
        "sunmaktadır."
    )
    add_para(
        "Dijital ve Yeşil Dönüşüm Destek Programları: Sanayi ve Teknoloji Bakanlığı, imalat "
        "sanayinde dijital olgunluk seviyesini artırmak amacıyla iki kapsamlı destek mekanizması "
        "işletmektedir. Dijital Dönüşüm Programı; 5 yıl ve üzeri faaliyet gösteren imalat KOBİ "
        "ve büyük işletmelerinin makine, yazılım ve danışmanlık giderlerini hibe ile "
        "desteklemektedir. KOSGEB bünyesinde 300 milyon Euro kaynak dijital dönüşüm projelerine "
        "ayrılmış olup hibe oranı %80'e, destek tutarı 1,5 milyon TL'ye çıkabilmektedir. Yeşil "
        "Dönüşüm Programı ise Dünya Bankası'nın 1 milyar dolarlık finansmanı çerçevesinde kaynak "
        "verimliliği, atık azaltma ve güneş enerjisi yatırımlarını desteklemekte; TÜBİTAK-TEYDEB'in "
        "1832 ve 1833 çağrı programları aracılığıyla 231 proje değerlendirilmektedir."
    )
    add_para(
        "IMF'nin yapay zeka patlamasının 1990'ların dot-com patlamasına benzediği uyarısı dikkate "
        "alınarak, AI yatırımlarında verimlilik odaklı bir yaklaşım benimsenmelidir. Teknoloji "
        "benimseme hızı sektörler arasında eşitsiz olduğundan, dijital beceri geliştirme "
        "programlarının sektörel ihtiyaçlara göre farklılaştırılması gerekmektedir (IMF, 2025a)."
    )

    doc.add_heading("7.6. Beceri Açığı, İşlerin Dönüşümü ve Aktif İşgücü Programları", level=2)

    add_para(
        "İstihdam teşvikleri tek başına yeterli değildir. Paralel olarak işgücünün dönüşen iş "
        "yapısına uyumunun sağlanması gerekmektedir. Bu iki politika aracı birlikte tasarlanmadığı "
        "sürece teşvikler kısa vadeli etki yaratır; sürdürülebilir verimlilik artışı sağlanamaz."
    )

    doc.add_heading("İşlerin Dönüşümü: Türkiye'de Hangi Meslekler Değişiyor?", level=3)
    add_para(
        "WEF Mesleklerin Geleceği 2025 Raporu, 2030'a kadar mevcut iş becerilerinin %39'unun "
        "geçerliliğini yitireceğini ve işverenlerin %41'inin yapay zeka otomasyonu nedeniyle "
        "işgücü azaltmayı planladığını ortaya koymaktadır (WEF, 2025). Bu dönüşüm Türkiye'de "
        "sektörel bazda somut biçimlerde kendini göstermektedir."
    )

    job_transformation = [
        ("İmalat Sanayi — Otomasyon Baskısı: ",
         "CNC tezgâhları, robotik kaynak sistemleri ve görüntü işleme kalite kontrol sistemleri, "
         "rutin montaj ve basit işleme görevlerini hızla otomatize etmektedir. Yok olan meslekler: "
         "Manuel montaj operatörü, basit kalite kontrol gözetmeni. Dönüşen meslekler: CNC "
         "programcısı, robotik sistem bakım teknisyeni, endüstriyel IoT operatörü. Bu dönüşüm "
         "için yeni gereken beceriler; PLC programlama, sensör veri yorumlama ve otomasyon "
         "bakımıdır. Türkiye'nin mesleki eğitim müfredatlarının bu yetkinlikleri kapsayacak "
         "şekilde güncellenmesi zorunludur."),
        ("Çağrı Merkezi ve İdari Hizmetler — Yapay Zeka Yerinden Edişi: ",
         "Doğal dil işleme ve büyük dil modelleri, standart müşteri hizmetleri, randevu ve "
         "şikâyet yönetimi görevlerinin büyük bölümünü otomatize edebilecek kapasitededir. "
         "Yok olan meslekler: Rutin müşteri temsilcisi, veri giriş operatörü. Dönüşen meslekler: "
         "Yapay zeka eğiticisi, müşteri deneyimi tasarımcısı, karmaşık süreç yöneticisi. "
         "Türkiye'deki 1,4 milyon İdari Hizmetler çalışanının önemli bir bölümünü bu risk "
         "yakından ilgilendirmektedir."),
        ("Muhasebe ve Finans — Dijital Muhasebe Araçları: ",
         "E-fatura, e-defter ve bulut muhasebe yazılımları, basit muhasebe tutma ve kayıt "
         "görevlerini minimize etmiştir. Yok olan meslekler: Basit defter tutucusu, ödeme "
         "takip görevlisi. Dönüşen meslekler: Finansal analiz uzmanı, vergi danışmanı, "
         "CFO yardımcısı. Bu dönüşüm KOBİ'lerde çalışan düşük nitelikli muhasebe personelini "
         "doğrudan etkilemektedir."),
        ("Sağlık — Dijital Sağlık ve Tele-tıp Dönüşümü: ",
         "MHRS, tele-tıp ve yapay zeka destekli görüntü analizi sistemleri yeni roller "
         "yaratmaktadır. Dönüşen meslekler: Uzaktan hasta takip koordinatörü, dijital sağlık "
         "veri analisti, tele-tıp hemşiresi. Kırsal bölgelerdeki sağlık çalışanları uzaktan "
         "hizmet sunum modellerine adapte olmak durumundadır."),
        ("Lojistik ve Ulaştırma — Otonom Araçlar ve Akıllı Depo: ",
         "Otonom araç teknolojisi ve otomatik depo sistemleri (AGV, robotik paketleme) "
         "geleneksel sürücü ve depo görevlerini dönüştürmektedir. Yok olan meslekler: "
         "Manuel depo sayım görevlisi, basit rota kurye. Dönüşen meslekler: Otonom araç "
         "filoya yöneticisi, akıllı depo sistem operatörü, son mil lojistik koordinatörü."),
    ]
    for title, desc in job_transformation:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(desc)

    doc.add_heading("Türkiye'de İşgücü Uyum Politikası Ne Yapmalı?", level=3)
    add_para(
        "TÜSİAD'ın (2023) verilerine göre Türk gençlerinin üçte biri NEET (ne eğitimde ne "
        "istihdamda) konumundadır. Bu yapısal sorun, işgücü uyum kapasitesini en baştan "
        "kısıtlamaktadır. Card, Kluve ve Weber'in (2018) 200+ ALMP programını kapsayan "
        "meta-analizi, mesleki eğitim programlarının orta vadede en kalıcı istihdam ve kazanç "
        "artışını sağladığını ortaya koymaktadır. OECD ülkeleri GSYİH'nin ortalama %0,62'sini "
        "ALMP'ye ayırırken Türkiye bu benchmarkın belirgin altında kalmaktadır."
    )
    add_para(
        "İŞKUR'un mevcut programları —mesleki eğitim kursları, işbaşı eğitimi, girişimcilik "
        "desteği, istihdam garantili kurslar— değerli ancak yetersiz kalan araçlardır. Dünya "
        "Bankası'nın (2023) İŞKUR değerlendirmesi, gençler için işbaşı eğitim programlarının "
        "istihdam ve kazanç artışı sağladığını teyit etmektedir. Ancak programların ölçeği, "
        "sektörel uyumu ve işverenle koordinasyonu güçlendirilmesi gerekmektedir."
    )
    add_para(
        "Sektörel beceri dönüşümü için politika önerileri: (i) İmalat sanayi SGK teşviki ile "
        "eş zamanlı olarak CNC operatörlüğü, robotik bakım ve endüstriyel IoT kursları "
        "açılmalı; (ii) İdari Hizmetler teşviki ile birlikte müşteri deneyimi ve yapay zeka "
        "eğitimi entegre edilmeli; (iii) Bilgi-İletişim teşvikine koşut yazılım, veri bilimi "
        "ve siber güvenlik bootcamp'leri desteklenmeli; (iv) Sağlık sektörü teşvikiyle "
        "birlikte tele-tıp hemşiresi ve dijital sağlık koordinatörü yetiştirme programları "
        "hayata geçirilmeli; (v) Her teşvik başvurusuna sektörel eğitim planı eklenmesi şartı "
        "getirilmeli, eğitimi tamamlayan çalışanlar için ek SGK indirimi tanınmalıdır. "
        "Tüm bu programlar İŞKUR veri tabanı aracılığıyla izlenmeli ve Card et al. (2018) "
        "meta-analizinin önerdiği bağımsız etki değerlendirmesine tabi tutulmalıdır."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. SONUÇ VE DEĞERLENDİRME
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. SONUÇ VE DEĞERLENDİRME", level=1)

    add_para(
        "Bu rapor, Türkiye'nin sektörel işgücü yapısını küresel ekonomik gelişmeler ve uluslararası "
        "kurum önerileri perspektifinden analiz etmiştir. Temel bulgular şu şekilde özetlenebilir:"
    )

    conclusions = [
        "Türkiye ekonomisinde 7 sektör (İmalat, İnşaat, İdari Destek, Konaklama, Sağlık, Mesleki "
        "Faaliyetler, Eğitim) hem yüksek istihdam kapasitesine sahip hem de yüksek işgücü maliyeti "
        "taşımaktadır. Bu sektörler toplam istihdamın yaklaşık %70'ini oluşturmakta ve teşvik "
        "politikalarının birincil hedef grubudur. Kamu Yönetimi ve Savunma, GKD ölçüm metodolojisi "
        "nedeniyle teşvik kapsamı dışında tutulmuştur.",

        "AB CBAM'ın 2026'da yürürlüğe girmesi, Türkiye imalat sanayinin yeşil dönüşümünü acil bir "
        "öncelik haline getirmiştir. Yılda 8 milyar Euro'luk ihracatın korunması için karbon "
        "yoğunluğunun düşürülmesi, bu süreçte istihdamın desteklenmesi kritiktir.",

        "OECD'nin kadın işgücü katılımı, hizmetler sektörü düzenlemeleri ve kayıt dışılık "
        "konularındaki önerileri, teşvik politikalarının tasarımında gözetilmelidir. Okul öncesi "
        "eğitime GSYH'nin %0,8'inin ayrılması kişi başı geliri %6 artırabilir.",

        "Dijital dönüşüm ve yapay zeka alanında Bilgi ve İletişim sektörünün stratejik olarak "
        "desteklenmesi, Türkiye'nin küresel değer zincirlerinde yukarı doğru hareketini "
        "hızlandıracaktır.",

        "Teşvik politikalarında ölü ağırlık kaybı, ikame etkisi ve mali sürdürülebilirlik riskleri "
        "dikkatlice yönetilmeli; net istihdam artışı koşulu, kademeli çıkış mekanizması ve "
        "yıllık maliyet-fayda analizi zorunlu tutulmalıdır.",
    ]
    for c in conclusions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c)
        p.paragraph_format.first_line_indent = Cm(0)

    add_para(
        "Sonuç olarak, Türkiye'nin istihdam teşviki politikaları tek boyutlu değil, sektörel, "
        "bölgesel, cinsiyet ve nitelik düzeyine göre farklılaştırılmış; küresel megatrendlerle "
        "(yeşil dönüşüm, dijitalleşme, AI) uyumlu ve uluslararası kurum önerileriyle tutarlı "
        "bir şekilde tasarlanmalıdır."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # KAYNAKÇA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("KAYNAKÇA", level=1)

    references = [
        "European Commission (2025). Carbon Border Adjustment Mechanism (CBAM). "
        "https://taxation-customs.ec.europa.eu/carbon-border-adjustment-mechanism_en",

        "France Stratégie (2020). Les Effets du CICE sur l'Emploi, l'Investissement et "
        "la Compétitivité des Entreprises. Comité de Suivi du CICE.",

        "GOV.UK (2025). Changes to the Class 1 National Insurance Contributions Secondary "
        "Threshold. https://www.gov.uk/government/publications/changes-to-the-class-1-national"
        "-insurance-contributions-secondary-threshold",

        "Hartmann, M., Kluve, J. ve Schmitz, S. (2013). Identifying Workers at Risk of Deadweight "
        "Loss and Substitution Effects. IAB Discussion Paper No. 05/2013. Institute for Employment "
        "Research. https://doku.iab.de/discussionpapers/2013/dp0513.pdf",

        "ICAP (2025). Türkiye Adopts Landmark Climate Law, Paving the Way for National ETS. "
        "International Carbon Action Partnership. "
        "https://icapcarbonaction.com/en/news/turkiye-adopts-landmark-climate-law-paving-way-national-ets",

        "IISD (2025). EU Carbon Border Adjustment Mechanism Is Set to Get Bigger: Implications "
        "for Trade and Industrial Value Chains. International Institute for Sustainable Development. "
        "https://www.iisd.org/articles/explainer/eu-carbon-border-adjustment-mechanism-bigger-trade-implications",

        "ILO (2025). World Employment and Social Outlook: Trends 2025. International Labour "
        "Organization. "
        "https://www.ilo.org/publications/flagship-reports/world-employment-and-social-outlook-trends-2025",

        "IZA (2023). Can Hiring Subsidies Benefit the Unemployed? IZA World of Labor. "
        "https://wol.iza.org/articles/can-hiring-subsidies-benefit-unemployed/long",

        "IMF (2025a). World Economic Outlook, October 2025: Global Economy in Flux, Prospects "
        "Remain Dim. International Monetary Fund. "
        "https://www.imf.org/en/publications/weo/issues/2025/10/14/world-economic-outlook-october-2025",

        "IMF (2025b). World Economic Outlook Update, July 2025: Global Economy: Tenuous Resilience "
        "amid Persistent Uncertainty. International Monetary Fund. "
        "https://www.imf.org/en/publications/weo/issues/2025/07/29/world-economic-outlook-update-july-2025",

        "IMF (2026). World Economic Outlook Update, January 2026: Global Economy: Steady amid "
        "Divergent Forces. International Monetary Fund. "
        "https://www.imf.org/en/publications/weo/issues/2026/01/19/world-economic-outlook-update-january-2026",

        "MOEL (2024). Employment Policy Overview. Republic of Korea Ministry of Employment "
        "and Labor. https://www.moel.go.kr/english/policy/employment.do",

        "OECD (2024b). Breaking the Vicious Circles of Informal Employment and Low-Paying Work. "
        "OECD Publishing. "
        "https://www.oecd.org/content/dam/oecd/en/publications/reports/2024/01/breaking-the-vicious-circles"
        "-of-informal-employment-and-low-paying-work_040b6f24/f95c5a74-en.pdf",

        "OECD (2025). OECD Economic Surveys: Türkiye 2025. Organisation for Economic Co-operation "
        "and Development. "
        "https://www.oecd.org/en/publications/oecd-economic-surveys-turkiye-2025_d01c660f-en.html",

        "OECD (2025b). OECD Compendium of Productivity Indicators 2025. Organisation for Economic "
        "Co-operation and Development. "
        "https://www.oecd.org/en/publications/oecd-compendium-of-productivity-indicators-2025_b024d9e1-en.html",

        "OECD (2025c). Türkiye: OECD Economic Outlook, Volume 2025 Issue 2. "
        "https://www.oecd.org/en/publications/2025/12/oecd-economic-outlook-volume-2025-issue-2_413f7d0a/full-report/turkiye_59622ed2.html",

        "OECD (2025d). Taxing Wages 2025: Tax Burden on Labour Income in OECD Countries. "
        "Organisation for Economic Co-operation and Development. "
        "https://www.oecd.org/en/publications/2025/04/taxing-wages-2025_20d1a01d.html",

        "Tax Foundation (2025). Kurzarbeit: Germany's Short-Work Subsidy Scheme. "
        "https://taxfoundation.org/blog/kurzarbeit-germany-short-work-subsidy-scheme/",

        "URSSAF (2025). Réduction Générale des Cotisations Patronales. "
        "https://www.urssaf.fr/accueil/employeur/beneficier-exonerations/reduction-generale-cotisation.html",

        "SGK (2024). Yıllık Bölüm 1: Sigortalı ve İş Yeri İstatistikleri, 2024. Sosyal Güvenlik "
        "Kurumu.",

        "T.C. Cumhurbaşkanlığı Strateji ve Bütçe Başkanlığı (2025). Orta Vadeli Program (2026-2028). "
        "Resmî Gazete.",

        "T.C. Cumhurbaşkanlığı (2024). Ulusal Yapay Zeka Stratejisi 2021-2025 (2024 Güncelleme). "
        "https://regulations.ai/regulations/turkey-ai-strategy-2021-2025",

        "TÜİK (2024). İktisadi Faaliyet Kollarına Göre Gelir Yöntemiyle Gayrisafi Yurtiçi Hasıla "
        "(Cari Fiyatlarla), Tablo I.2.14. Türkiye İstatistik Kurumu.",

        "UNIDO (2024). Industrial Development Report 2024. United Nations Industrial Development "
        "Organization. https://www.unido.org/sites/default/files/unido-publications/2023-11/IDR24-OVERVIEW_1.pdf",

        "UNIDO (2025). International Yearbook of Industrial Statistics 2025. United Nations Industrial "
        "Development Organization. "
        "https://www.unido.org/sites/default/files/unido-publications/2025-11/UNIDO%20Industrial%20StatisticsYearbook%202025.pdf",

        "VoxDev (2024). Can Temporary Wage Incentives Increase Formal Employment? Experimental "
        "Evidence from Mexico. "
        "https://voxdev.org/topic/labour-markets/can-temporary-wage-incentives-increase-formal-employment-experimental-evidence",

        "WEF (2025). The Impact of the EU's CBAM on Business and the Carbon-Pricing Landscape. "
        "World Economic Forum. "
        "https://www.weforum.org/stories/2025/12/eu-cbam-impact-business-carbon-pricing-landscape/",

        "World Bank (2025). Türkiye Country Economic Memorandum: Jobs for Prosperity. "
        "https://documents.worldbank.org/en/publication/documents-reports/documentdetail/099091725075068918",

        "World Bank (2025b). Türkiye Macro Poverty Outlook. "
        "https://thedocs.worldbank.org/en/doc/d5f32ef28464d01f195827b7e020a3e8-0500022021/related/mpo-tur.pdf",

        "Brookings Institution (2015). America's Advanced Industries: New Trends. "
        "Brookings Institution Press.",

        "Card, D., Kluve, J. ve Weber, A. (2018). What Works? A Meta Analysis of Recent Active "
        "Labor Market Program Evaluations. Journal of the European Economic Association, 16(3), "
        "894-931.",

        "IZA (2016). Do Employment Subsidies Work? Evidence from Regionally Targeted Subsidies "
        "in Turkey. IZA Discussion Paper No. 9993. Institute of Labor Economics. "
        "https://docs.iza.org/dp9993.pdf",

        "Moretti, E. (2010). Local Multipliers. American Economic Review: Papers and Proceedings, "
        "100(2), 373-377.",

        "NAM — National Association of Manufacturers (2025). Manufacturing Economic Impact. "
        "https://nam.org/manufacturing-economic-impact/",

        "OECD (2015). OECD Product Market Regulation Indicators for Türkiye. Organisation for "
        "Economic Co-operation and Development. "
        "https://stats.oecd.org/Index.aspx?DataSetCode=PMR",

        "TEPAV (2024). Türkiye Ekonomi Görünümü. Ekonomi Politikaları Araştırma Vakfı. "
        "https://www.tepav.org.tr",

        "TÜSİAD (2023). Türkiye'de Dijital Dönüşümün İstihdam ve Beceriler Üzerindeki Etkisi. "
        "Türk Sanayicileri ve İş İnsanları Derneği.",

        "WEF (2025). Future of Jobs Report 2025. World Economic Forum. "
        "https://reports.weforum.org/docs/WEF_Future_of_Jobs_Report_2025.pdf",

        "World Bank (2023). Türkiye: Evaluating the Impact of İŞKUR Vocational Training Programs. "
        "World Bank Group.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.hanging_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        run.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════════════════
    # KAYDET
    # ══════════════════════════════════════════════════════════════════════════
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    ana = os.path.join(os.path.dirname(__file__), "..",
                       "SEKTÖR BAZLI İŞGÜCÜ EKONOMİK GÖSTERGESİ.xlsx")
    sgk_dir = os.path.join(os.path.dirname(__file__), "..", "sgk_veriler")
    sgk = None
    for f in os.listdir(sgk_dir):
        if "BÖLÜM 1" in f and f.endswith(".xlsx"):
            sgk = os.path.join(sgk_dir, f)
            break

    data = load_all_data(ana, sgk)
    output = os.path.join(os.path.dirname(__file__), "..",
                          "KURESEL_GELISMELER_ISIGINDA_SEKTOREL_ANALIZ_RAPORU.docx")
    create_academic_report(data, output)
    print(f"Rapor oluşturuldu: {output}")
